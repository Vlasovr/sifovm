from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
SCHEMES = ASSETS / "schemes"
WAVES = ASSETS / "waveforms"

STUDENT = "Власов Р.Е."
GROUP = "250541"
CITY_YEAR = "Минск 2026"
DISCIPLINE = "Структурная и функциональная организация вычислительных машин"


@dataclass
class LabReport:
    num: str
    title: str
    topic: str
    folder: str
    file_name: str
    qpf: str
    top: str
    scheme: str
    wave: str
    toc_pages: list[tuple[str, str]]
    variant_rows: list[list[str]]
    theory: list[str]
    theory_table: list[list[str]]
    task_rows: list[list[str]]
    signal_rows: list[list[str]]
    algorithm_rows: list[list[str]]
    module_rows: list[list[str]]
    implementation: list[str]
    test_rows: list[list[str]]
    resource_rows: list[list[str]]
    conclusion: list[str]


def set_run_font(run, size: float = 14, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(12 if name == "Heading 1" else 6)


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, 12)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def para(
    doc: Document,
    text: str = "",
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_indent: bool = True,
    size: float = 14,
    bold: bool = False,
    italic: bool = False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, italic=italic)
    return paragraph


def section_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    set_run_font(run, 14, bold=True)
    return paragraph


def subheading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, 14, bold=True)
    return paragraph


def shade_cell(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 11, align=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = align or (WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 16 else WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    font_size: float = 11,
    header_fill: str = "D9EAF7",
    align: WD_TABLE_ALIGNMENT = WD_TABLE_ALIGNMENT.CENTER,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = align
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=font_size)
        shade_cell(table.rows[0].cells[index], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_toc_table(doc: Document, entries: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for title, page in entries:
        cells = table.add_row().cells
        set_cell_text(cells[0], title, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[1], page, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()


def add_picture(doc: Document, path: Path, caption: str, *, width_cm: float = 15.4):
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_p = para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, size=12, italic=True)
    caption_p.paragraph_format.space_after = Pt(6)


def title_page(doc: Document, lab: LabReport):
    title_lines = [
        "Министерство образования Республики Беларусь",
        "Учреждение образования",
        "Белорусский государственный университет",
        "информатики и радиоэлектроники",
        "Факультет компьютерных систем и сетей",
        "Кафедра электронных вычислительных машин",
        f"Дисциплина: {DISCIPLINE}",
    ]
    for line in title_lines:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, size=12)

    for _ in range(6):
        doc.add_paragraph()

    para(doc, "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True)
    para(doc, "на тему", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    para(doc, lab.topic, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    for _ in range(8):
        doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    left = [
        f"Выполнил\nстудент гр. {GROUP}",
        "Проверил",
    ]
    right = [
        STUDENT,
        "__________________",
    ]
    for row_index in range(2):
        set_cell_text(table.rows[row_index].cells[0], left[row_index], size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[row_index].cells[1], right[row_index], size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)

    for _ in range(4):
        doc.add_paragraph()

    para(doc, CITY_YEAR, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, size=12)
    doc.add_page_break()


def add_common_opening(doc: Document, lab: LabReport):
    section_heading(doc, "СОДЕРЖАНИЕ")
    add_toc_table(doc, lab.toc_pages)

    section_heading(doc, "1 Цель работы")
    para(
        doc,
        f"Целью лабораторной работы является разработка, описание, функциональная проверка "
        f"и синтез устройства «{lab.title}» для индивидуального варианта микро-ЭВМ. "
        "Работа выполняется так, чтобы разработанный узел мог быть использован как самостоятельный "
        "проект Quartus и как часть последующей общей структуры процессора.",
    )
    para(
        doc,
        "В ходе выполнения необходимо определить назначение входных и выходных сигналов, "
        "разработать алгоритм функционирования устройства, подготовить его структурное описание, "
        "подготовить проект Quartus II 9.1, выполнить моделирование и подтвердить правильность "
        "работы контрольными наборами.",
    )
    add_table(doc, ["Этап", "Содержание работы"], lab.task_rows, font_size=10.5)
    doc.add_page_break()

    section_heading(doc, "2 Исходные данные к работе")
    para(
        doc,
        "Лабораторная работа выполняется в среде Quartus II 9.1. Разрядность основных трактов данных "
        "принята равной 16 бит, "
        "что соответствует принятому варианту курсового проекта.",
    )
    para(
        doc,
        "Верхний уровень проекта оформлен так, чтобы внешние линии имели развернутые имена. "
        "По ним сразу определяется назначение команды, операндов, флагов, линий памяти "
        "и диагностических сигналов, что упрощает чтение схемы и проверку временных диаграмм.",
    )
    add_table(doc, ["Параметр", "Значение"], lab.variant_rows, font_size=10.5)
    doc.add_page_break()


def add_theory(doc: Document, lab: LabReport):
    section_heading(doc, "3 Теоретические сведения")
    for paragraph in lab.theory:
        para(doc, paragraph)
    add_table(doc, ["Понятие", "Назначение в разрабатываемом устройстве"], lab.theory_table, font_size=10.0)
    doc.add_page_break()


def add_execution(doc: Document, lab: LabReport):
    section_heading(doc, "4 Выполнение работы")

    subheading(doc, "4.1 Структура разрабатываемого устройства")
    para(
        doc,
        "Структура проекта выбрана модульной: каждый функциональный узел вынесен в отдельный блок, "
        "а верхний модуль соединяет эти узлы сигналами данных и управления. Такой способ удобен для проверки, "
        "так как отдельные блоки можно независимо анализировать и проверять моделированием.",
    )
    add_picture(doc, SCHEMES / lab.scheme, f"Рисунок 4.1 - Структура устройства «{lab.title}»", width_cm=15.2)

    subheading(doc, "4.2 Входные и выходные сигналы")
    para(
        doc,
        "В таблице приведены основные точки наблюдения верхнего уровня. Эти же имена используются "
        "на схемах и временных диаграммах при проверке работы устройства.",
    )
    add_table(doc, ["Сигнал", "Направление", "Разрядность", "Назначение"], lab.signal_rows, font_size=9.0)

    subheading(doc, "4.3 Алгоритм работы")
    para(
        doc,
        "Алгоритм работы устройства задает последовательность преобразования входных сигналов в выходные "
        "и определяет условия перехода между режимами. Для синхронных блоков изменение регистров выполняется "
        "по фронту тактового сигнала, для комбинационных блоков результат формируется после изменения входов.",
    )
    add_table(doc, ["Шаг", "Действие", "Результат"], lab.algorithm_rows, font_size=9.5)

    subheading(doc, "4.4 Реализация функциональных блоков")
    for paragraph in lab.implementation:
        para(doc, paragraph)
    add_table(doc, ["Блок проекта", "Назначение"], lab.module_rows, font_size=9.5)
    para(
        doc,
        f"Файл проекта Quartus: {lab.folder}/quartus/{lab.qpf}. "
        "После открытия проекта необходимо выполнить Processing -> Start Compilation и проверить "
        "схемное представление верхнего уровня устройства.",
    )

    subheading(doc, "4.5 Функциональное моделирование")
    para(
        doc,
        "Функциональная проверка выполнена testbench-моделью. В testbench подаются контрольные наборы, "
        "соответствующие рабочим и граничным режимам устройства. Проверка результата выполняется операторами "
        "assert, поэтому ошибка в данных, флагах или протоколе приводит к останову моделирования.",
    )
    add_table(doc, ["Проверка", "Условие", "Ожидаемый результат"], lab.test_rows, font_size=9.5)
    add_picture(doc, WAVES / lab.wave, f"Рисунок 4.2 - Контрольная временная диаграмма для ЛР{lab.num}", width_cm=14.4)

    subheading(doc, "4.6 Результаты синтеза")
    para(
        doc,
        "Компиляция проекта выполняется в Quartus II 9.1. Описание устройства не использует вручную "
        "размещенные вентильные примитивы, поэтому схема синтезируется средствами Quartus на выбранную "
        "элементную базу. Основные показатели компиляции приведены в таблице.",
    )
    add_table(doc, ["Показатель", "Значение"], lab.resource_rows, font_size=10.0)
    doc.add_page_break()


def add_conclusion(doc: Document, lab: LabReport):
    section_heading(doc, "5 Вывод")
    for paragraph in lab.conclusion:
        para(doc, paragraph)


def neutralize_text(text: str) -> str:
    replacements = [
        ("VHDL-entity", "функциональный блок"),
        ("VHDL-описание", "описание устройства"),
        ("VHDL-блоки", "функциональные блоки"),
        ("VHDL-модулей", "функциональных модулей"),
        ("на языке VHDL", "в виде структурного описания"),
        ("языке VHDL", "структурном описании"),
        ("VHDL", "структурное описание"),
        ("RTL Viewer", "схемное представление проекта"),
        ("RTL-схеме", "схеме проекта"),
        ("testbench-моделью", "тестовой моделью"),
        ("testbench", "тестовая модель"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def neutralize_rows(rows: list[list[str]]) -> list[list[str]]:
    return [[neutralize_text(cell) for cell in row] for row in rows]


def apply_common_neutral_style(lab: LabReport) -> LabReport:
    lab.toc_pages = [(neutralize_text(a), neutralize_text(b)) for a, b in lab.toc_pages]
    lab.variant_rows = neutralize_rows(lab.variant_rows)
    lab.theory = [neutralize_text(p) for p in lab.theory]
    lab.theory_table = neutralize_rows(lab.theory_table)
    lab.task_rows = neutralize_rows(lab.task_rows)
    lab.signal_rows = neutralize_rows(lab.signal_rows)
    lab.algorithm_rows = neutralize_rows(lab.algorithm_rows)
    lab.module_rows = neutralize_rows(lab.module_rows)
    lab.implementation = [neutralize_text(p) for p in lab.implementation]
    lab.test_rows = neutralize_rows(lab.test_rows)
    lab.resource_rows = neutralize_rows(lab.resource_rows)
    lab.conclusion = [neutralize_text(p) for p in lab.conclusion]
    return lab


def prepare_variant4_lab(lab: LabReport) -> LabReport:
    if lab.num == "5":
        lab.variant_rows = [
            ["Разрядность операндов", "16 бит"],
            ["Команда сдвига", "SRA - арифметический сдвиг вправо"],
            ["Арифметическая команда", "CMP, прямая регистровая адресация"],
            ["Логическая команда", "NOR"],
            ["Формируемые признаки", "Z, S, C, O"],
            ["Коды операций", "CMP=03h, NOR=04h, SRA=05h"],
        ]
        lab.theory = [
            "Арифметико-логическое устройство является исполнительным узлом процессора. Оно принимает код операции и два 16-разрядных операнда, выполняет выбранное действие и формирует признаки результата. Для варианта 4 реализуются арифметический сдвиг вправо SRA, сравнение CMP с прямой регистровой адресацией и логическая операция NOR.",
            "Команда CMP выполняет вычитание второго операнда из первого. Полученная разность используется для формирования признаков: Z устанавливается при равенстве операндов, S показывает знак разности, C фиксирует заем при A<B, O показывает переполнение при знаковом сравнении.",
            "Операция NOR является инверсией поразрядного OR. Она удобна для проверки логической части АЛУ и для формирования нулевого результата при совпадающих единичных операндах.",
            "Операция SRA выполняет арифметический сдвиг вправо. Старший разряд результата повторяет старший разряд исходного слова, поэтому знак числа сохраняется. Младший вытолкнутый бит помещается в признак C.",
        ]
        lab.theory_table = [
            ["CMP", "Y=A-B; Z=1 при равенстве, S равен старшему биту разности, C показывает заем, O показывает знаковое переполнение."],
            ["NOR", "Y=NOT(A OR B); результат формируется по всем 16 разрядам."],
            ["SRA", "Y(15)=A(15), Y(14..0)=A(15..1), C=A(0); знак сохраняется."],
            ["Флаги", "Z и S формируются по результату, C и O поступают из выбранной операции."],
        ]
        lab.signal_rows = [
            ["cmd_opcode_alu_operation_i", "in", "8", "Код команды АЛУ: CMP, NOR или SRA."],
            ["operand_a_from_register_i", "in", "16", "Первый операнд A из регистра общего назначения."],
            ["operand_b_from_memory_or_reg_i", "in", "16", "Второй операнд B для CMP и NOR."],
            ["result_y_to_register_file_o", "out", "16", "Результат выбранной операции."],
            ["flag_z_zero_result_o", "out", "1", "Признак нулевого результата."],
            ["flag_s_negative_result_o", "out", "1", "Признак отрицательного результата."],
            ["flag_c_carry_or_shift_out_o", "out", "1", "Заем при CMP или вытолкнутый бит при SRA."],
            ["flag_o_overflow_o", "out", "1", "Признак знакового переполнения при CMP."],
            ["ctrl_selected_alu_operation_o", "out", "3", "Диагностический код выбранной ветви АЛУ."],
            ["ctrl_second_operand_is_used_o", "out", "1", "Признак использования второго операнда."],
            ["ctrl_write_flags_register_o", "out", "1", "Разрешение записи регистра признаков."],
        ]
        lab.algorithm_rows = [
            ["1", "Блок управления сравнивает opcode с кодами CMP, NOR и SRA.", "Формируется внутренний код операции и разрешение записи признаков."],
            ["2", "Ветви CMP, NOR и SRA параллельно вычисляют возможные результаты.", "Задержка определяется выбором готового результата мультиплексором."],
            ["3", "Мультиплексор выбирает результат по внутреннему коду операции.", "На выход Y поступает разность, логический результат или сдвинутое слово."],
            ["4", "Формирователь признаков анализирует выбранный результат и служебные признаки операции.", "Формируются Z, S, C и O."],
        ]
        lab.module_rows = [
            ["Общие параметры варианта", "Разрядность данных, коды команд и внутренние коды операций."],
            ["Блок управления АЛУ", "Дешифрация команды и выбор ветви CMP, NOR или SRA."],
            ["Блок сравнения CMP", "Вычитание A-B и формирование признаков заема и переполнения."],
            ["Блок операции NOR", "Поразрядная логическая операция NOT(A OR B)."],
            ["Блок сдвига SRA", "Арифметический сдвиг вправо с сохранением знака."],
            ["Формирователь признаков", "Выработка Z, S, C и O."],
            ["Верхний уровень АЛУ", "Соединение функциональных блоков и внешних линий."],
            ["Тестовая модель", "Проверка CMP, NOR, SRA и признаков результата."],
        ]
        lab.implementation = [
            "Внутри АЛУ использована параллельная структура: все основные ветви получают операнды одновременно, а затем мультиплексор выбирает результат по коду операции. Такой способ соответствует аппаратной реализации и позволяет явно показать блок сравнения, блок логической операции и блок сдвига.",
            "Для CMP выполняется вычитание A-B с расширением разрядности на один бит. Это позволяет корректно определить заем. Переполнение определяется по знакам операндов и знаку результата. Для SRA переносом считается младший бит исходного слова, а для NOR признаки C и O сбрасываются.",
        ]
        lab.test_rows = [
            ["CMP", "A=0010h, B=0001h", "Y=000Fh, Z=0, S=0, C=0"],
            ["CMP равенство", "A=1234h, B=1234h", "Y=0000h, Z=1"],
            ["CMP с заемом", "A=0001h, B=0002h", "Y=FFFFh, S=1, C=1"],
            ["NOR", "A=0FFFh, B=0FFFh", "Y=F000h, S=1"],
            ["SRA", "A=8001h", "Y=C000h, C=1, знак сохранен"],
            ["NOR нулевой результат", "A=FFFFh, B=FFFFh", "Y=0000h, Z=1"],
        ]
        lab.resource_rows = [
            ["Combinational ALUTs", "59"],
            ["Dedicated logic registers", "0"],
            ["Pins", "65"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ]
        lab.conclusion = [
            "В ходе лабораторной работы разработано арифметико-логическое устройство для варианта 4. Реализованы операции CMP, NOR и SRA, а также формирование признаков Z, S, C и O.",
            "Моделирование подтвердило корректность сравнения, логической операции и арифметического сдвига. Проект успешно компилируется в Quartus II 9.1 и может использоваться как исполнительный узел микро-ЭВМ.",
        ]

    elif lab.num == "6":
        lab.variant_rows = [
            ["Объем стека", "7 слов по 16 бит"],
            ["Направление роста", "Вверх, от младших адресов к старшим"],
            ["Указатель стека", "Указывает на последнюю занятую ячейку"],
            ["Пустой стек", "SP=7 как специальное состояние вне диапазона ячеек"],
            ["Полный стек", "SP=6"],
            ["Дополнительный режим", "Занесение результата CMP из АЛУ в стек"],
        ]
        lab.theory = [
            "Стековое запоминающее устройство работает по принципу LIFO: последним записанное слово считывается первым. В варианте 4 стек содержит семь 16-разрядных слов, растет вверх, а указатель SP показывает последнюю занятую ячейку.",
            "Поскольку при пустом стеке последней занятой ячейки нет, используется специальное значение SP=7. После первой операции PUSH данные записываются в ячейку 0, а SP становится равным 0. При заполнении всех семи ячеек SP принимает значение 6.",
            "При выполнении POP данные считываются из ячейки, на которую указывает SP, после чего указатель уменьшается. Если считывалась ячейка 0, стек становится пустым и SP возвращается в состояние 7.",
        ]
        lab.theory_table = [
            ["PUSH", "Если стек не полон, новое слово записывается в следующую старшую ячейку."],
            ["POP", "Если стек не пуст, считывается слово из последней занятой ячейки."],
            ["PUSH_ALU", "В стек помещается результат арифметической операции предыдущей лабораторной работы."],
            ["overflow", "Устанавливается при попытке записи в полный стек."],
            ["underflow", "Устанавливается при попытке чтения из пустого стека."],
        ]
        lab.algorithm_rows = [
            ["1", "После сброса память стека очищается, SP=7, empty=1.", "Стек готов к первой записи."],
            ["2", "При PUSH выбирается источник данных: РОН или результат АЛУ.", "Данные записываются в следующую ячейку вверх."],
            ["3", "При POP считывается ячейка, на которую указывает SP.", "На выход передается последнее записанное слово."],
            ["4", "Контрольная логика проверяет границы SP.", "Формируются full, empty, overflow и underflow."],
        ]
        lab.module_rows = [
            ["Общие параметры варианта", "Глубина стека, разрядность слова и коды команд."],
            ["Блок управления стеком", "Дешифрация PUSH, POP и PUSH_ALU, выбор источника данных."],
            ["Регистровый массив стека", "Семь 16-разрядных ячеек, указатель SP и признаки состояния."],
            ["Верхний уровень стекового устройства", "Соединение управления, мультиплексора входа и массива стека."],
            ["Тестовая модель", "Проверка PUSH, POP, PUSH_ALU, overflow и underflow."],
        ]
        lab.implementation = [
            "Стек реализован как синхронный регистровый массив из семи 16-разрядных слов. Запись и чтение выполняются по фронту тактового сигнала. Направление роста вверх реализовано увеличением SP при успешной записи и уменьшением SP при успешном чтении.",
            "Значение SP=7 используется только как признак пустого стека. Это позволяет сохранить смысл варианта: при наличии данных SP всегда указывает на последнюю занятую ячейку в диапазоне 0..6.",
        ]
        lab.test_rows = [
            ["Сброс", "reset=1", "SP=7, empty=1, full=0"],
            ["Первый PUSH", "DIN=1111h", "SP=0, занята ячейка 0"],
            ["Второй PUSH и POP", "DIN=2222h, затем POP", "DOUT=2222h, SP=0"],
            ["PUSH_ALU", "DIN_ALU=000Fh", "Результат CMP помещается в стек"],
            ["Заполнение", "Всего семь занятых слов", "SP=6, full=1"],
            ["Границы", "PUSH в полный стек и POP из пустого", "overflow=1 и underflow=1"],
        ]
        lab.resource_rows = [
            ["Combinational ALUTs", "68"],
            ["Dedicated logic registers", "135"],
            ["Pins", "61"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ]
        lab.conclusion = [
            "В ходе работы разработано стековое запоминающее устройство варианта 4: глубина 7 слов, рост вверх, указатель на последнюю занятую ячейку.",
            "Проверка подтвердила правило LIFO, корректную работу PUSH/POP/PUSH_ALU и защиту от переполнения и чтения из пустого стека. Проект успешно компилируется в Quartus II 9.1.",
        ]

    elif lab.num == "7":
        lab.module_rows = [
            ["Ведущее устройство", "Формирование запроса, слова данных и признака обслуживания."],
            ["Центральный арбитр", "Выбор ведущего по фиксированному временному кванту."],
            ["Мультиплексор общей шины", "Передача слова выбранного ведущего на общую шину."],
            ["Ведомое устройство", "Синхронный прием данных по сигналу занятости шины."],
            ["Верхний уровень системы", "Соединение ведущих, арбитра, мультиплексора и ведомого."],
            ["Тестовая модель", "Проверка one-hot grant, соответствия REQ/GNT и передачи данных."],
        ]
        lab.resource_rows = [
            ["Combinational ALUTs", "34"],
            ["Dedicated logic registers", "19"],
            ["Pins", "49"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ]

    elif lab.num == "8":
        lab.variant_rows = [
            ["Вид кэш-памяти", "Прямое отображение"],
            ["Количество строк", "8"],
            ["Количество слов в строке", "1 слово, 16 бит"],
            ["Количество строк в наборе", "1"],
            ["Принцип замещения", "Не требуется: индекс однозначно выбирает строку"],
            ["Политика записи", "Write-through с обновлением основной памяти"],
        ]
        lab.theory = [
            "Кэш-память с прямым отображением разбивает адрес на тег и индекс. Индекс выбирает ровно одну строку кэша, а тег хранится вместе с данными и используется для проверки принадлежности строки текущему адресу.",
            "В варианте 4 используется восемь строк. Поэтому для выбора строки достаточно трех младших разрядов адреса A[2..0], а старшие разряды A[15..3] образуют тег. В каждой строке хранится одно 16-разрядное слово.",
            "При совпадении valid-бита и тега возникает попадание. При промахе выбранная индексом строка перезаписывается словом из основной памяти. Отдельный алгоритм выбора жертвы не нужен, потому что для каждого индекса существует только одна строка.",
        ]
        lab.theory_table = [
            ["TAG", "Старшая часть адреса, сравниваемая с тегом выбранной строки."],
            ["INDEX", "Три младших разряда адреса, выбирающие одну из восьми строк."],
            ["valid", "Признак того, что строка содержит актуальные данные."],
            ["hit", "Совпадение valid и tag."],
            ["miss", "Отсутствие нужного слова в выбранной строке, требуется обращение к основной памяти."],
        ]
        lab.signal_rows = [
            ["clock_i", "in", "1", "Тактовый сигнал контроллера кэша и модели памяти."],
            ["reset_i", "in", "1", "Сброс valid-битов, тегов, данных и автомата управления."],
            ["cpu_cache_request_i", "in", "1", "Запрос процессора на чтение или запись."],
            ["cpu_write_enable_i", "in", "1", "Тип операции: 0 - чтение, 1 - запись."],
            ["cpu_address_tag_set_i", "in", "16", "Адрес процессора: A15..A3 - тег, A2..A0 - индекс строки."],
            ["cpu_write_data_to_cache_i", "in", "16", "Данные процессора для операции записи."],
            ["cpu_read_data_from_cache_o", "out", "16", "Данные, возвращаемые процессору при чтении."],
            ["cpu_cache_response_ready_o", "out", "1", "Признак готовности ответа кэша."],
            ["cache_hit_signal_o", "out", "1", "Импульс попадания."],
            ["cache_miss_signal_o", "out", "1", "Импульс промаха."],
            ["memory_request_from_cache_debug_o", "out", "1", "Диагностический запрос к основной памяти."],
            ["memory_write_enable_debug_o", "out", "1", "Диагностический признак записи в основную память."],
            ["memory_address_debug_o", "out", "16", "Адрес, передаваемый из кэша в основную память."],
            ["memory_data_debug_o", "out", "16", "Данные на интерфейсе основной памяти."],
        ]
        lab.algorithm_rows = [
            ["1", "CPU подает запрос, адрес, тип операции и данные записи.", "Контроллер выделяет TAG и INDEX."],
            ["2", "По INDEX выбирается строка кэша.", "Проверяются valid и tag."],
            ["3", "При hit чтение завершается из кэша, при write-hit обновляются кэш и основная память.", "CPU получает ready."],
            ["4", "При miss выполняется обращение к основной памяти.", "Выбранная строка перезаписывается новым tag/data."],
        ]
        lab.module_rows = [
            ["Контроллер прямого отображения", "Массивы valid/tag/data, проверка hit/miss и автомат обращения к памяти."],
            ["Основная память", "Синхронная модель RAM для чтения и записи при промахах и write-through."],
            ["Верхний уровень кэш-памяти", "Соединение контроллера кэша и основной памяти."],
            ["Тестовая модель", "Проверка miss, hit, write-through и замещения строки с одинаковым индексом."],
        ]
        lab.implementation = [
            "В кэше заведены три массива: valid, tag и data. Индекс A[2..0] выбирает одну из восьми строк, после чего тег выбранной строки сравнивается с A[15..3].",
            "При чтении и попадании данные сразу выдаются процессору. При чтении и промахе контроллер запрашивает основную память, записывает полученное слово в выбранную строку и затем формирует ответ процессору.",
            "При записи используется политика write-through: слово записывается в кэш и одновременно передается в основную память. Если запись пришлась на адрес с новым тегом, выбранная строка также получает новый тег и valid-бит.",
        ]
        lab.test_rows = [
            ["Первое чтение", "READ 0010h", "miss=1, data=1110h, строка заполняется"],
            ["Повторное чтение", "READ 0010h", "hit=1, data=1110h"],
            ["Запись", "WRITE 0010h, data=ABCDh", "hit=1, кэш и память обновлены"],
            ["Конфликт индекса", "READ 0018h", "miss=1, строка с тем же index заменяется"],
            ["Повторное чтение старого адреса", "READ 0010h", "miss=1, data=ABCDh загружается из памяти"],
        ]
        lab.resource_rows = [
            ["Combinational ALUTs", "200"],
            ["Dedicated logic registers", "374"],
            ["Pins", "89"],
            ["Block memory bits", "4096"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ]
        lab.conclusion = [
            "В ходе работы разработана кэш-память с прямым отображением для варианта 4. Адрес разделен на TAG и INDEX, каждая строка содержит valid-бит, тег и 16-разрядное слово данных.",
            "Моделирование подтвердило корректную обработку попаданий, промахов, записи по write-through и замещения строки с одинаковым индексом. Проект успешно компилируется в Quartus II 9.1.",
        ]

    return apply_common_neutral_style(lab)


def report_data() -> list[LabReport]:
    common_task_rows = [
        ["1", "Изучить индивидуальный вариант и выделить параметры устройства."],
        ["2", "Разработать структурную схему, определить входные и выходные сигналы."],
        ["3", "Описать функциональные блоки на языке VHDL и собрать верхний уровень проекта."],
        ["4", "Подготовить проект Quartus II 9.1 и проверить корректность синтеза."],
        ["5", "Разработать testbench, выполнить моделирование и сравнить результат с ожидаемым."],
    ]

    lab5 = LabReport(
        num="5",
        title="Арифметико-логическое устройство",
        topic="«Арифметико-логическое устройство»",
        folder="lab5_alu",
        file_name="LR5_ALU_variant.docx",
        qpf="lab5_alu.qpf",
        top="lab5_alu_rtl_view_top",
        scheme="lab5_alu_scheme.png",
        wave="lab5_waveform.png",
        toc_pages=[
            ("1 Цель работы........................................................................", "3"),
            ("2 Исходные данные к работе....................................................", "4"),
            ("3 Теоретические сведения.......................................................", "5"),
            ("4 Выполнение работы..............................................................", "6"),
            ("5 Вывод....................................................................................", "10"),
        ],
        variant_rows=[
            ["Разрядность операндов", "16 бит"],
            ["Логические операции", "OR, NOR"],
            ["Операция сдвига", "SRA - арифметический сдвиг вправо"],
            ["Арифметическая операция", "INCS - инкремент при установленном флаге S"],
            ["Формируемые признаки", "Z, S, C, O"],
            ["Коды операций", "OR=03h, NOR=04h, SRA=05h, INCS=06h"],
        ],
        theory=[
            "Арифметико-логическое устройство является исполнительным узлом процессора. Оно принимает операнды из регистрового файла или памяти, выполняет выбранную операцию и возвращает результат в общий тракт данных. В варианте реализуются две логические операции, одна операция сдвига и одна арифметическая операция, зависящая от состояния регистра флагов.",
            "Операция OR выполняет поразрядное логическое сложение двух 16-разрядных слов. Для каждого разряда результата устанавливается единица, если хотя бы один из соответствующих разрядов операндов равен единице. Операция NOR является инверсией результата OR и используется для получения нулевого результата при совпадающих единичных операндах.",
            "Арифметический сдвиг вправо SRA сохраняет знаковый разряд. Старший бит результата получает значение старшего бита исходного слова, остальные разряды сдвигаются вправо, а младший бит исходного слова выводится в перенос C. Поэтому для отрицательных чисел в дополнительном коде знак после сдвига сохраняется.",
            "Операция INCS использует входной флаг S. Если флаг S равен единице, к операнду A прибавляется единица; если флаг S равен нулю, операнд передается без изменения. Такой вариант удобен при выполнении условного инкремента по признаку отрицательного результата предыдущей операции.",
            "После выбора результата формируются признаки Z, S, C и O. Признак Z устанавливается при нулевом результате, S повторяет старший бит результата, C фиксирует перенос или вытолкнутый при сдвиге бит, O используется для признака переполнения при арифметической операции.",
        ],
        theory_table=[
            ["OR", "Y = A OR B; используется для установки разрядов результата по единичным разрядам операндов."],
            ["NOR", "Y = NOT(A OR B); позволяет получить инверсный логический результат и проверить нулевой случай."],
            ["SRA", "Y(15)=A(15), Y(14..0)=A(15..1), C=A(0); сохраняется знак числа."],
            ["INCS", "Y=A+1 при FR.S=1, иначе Y=A; C и O формируются как для сложения с единицей."],
            ["Флаги", "Z и S формируются от результата, C и O приходят из операции сдвига или сложения."],
        ],
        task_rows=common_task_rows,
        signal_rows=[
            ["cmd_opcode_alu_operation_i", "in", "8", "Код операции, поступающий из регистра команды или блока управления."],
            ["operand_a_from_register_i", "in", "16", "Первый операнд A, обычно считанный из регистра общего назначения."],
            ["operand_b_from_memory_or_reg_i", "in", "16", "Второй операнд B для операций OR и NOR."],
            ["flag_s_from_flags_register_i", "in", "1", "Сохраненный признак S, используемый операцией INCS."],
            ["result_y_to_register_file_o", "out", "16", "Результат операции, передаваемый в регистровый файл или общий тракт данных."],
            ["flag_z_zero_result_o", "out", "1", "Признак нулевого результата."],
            ["flag_s_negative_result_o", "out", "1", "Признак отрицательного результата, равный старшему биту Y."],
            ["flag_c_carry_or_shift_out_o", "out", "1", "Перенос при INCS или вытолкнутый младший бит при SRA."],
            ["flag_o_overflow_o", "out", "1", "Признак переполнения для арифметической операции."],
            ["ctrl_selected_alu_operation_o", "out", "3", "Диагностический код выбранной внутренней операции АЛУ."],
            ["ctrl_second_operand_is_used_o", "out", "1", "Диагностический сигнал использования второго операнда."],
            ["ctrl_write_flags_register_o", "out", "1", "Разрешение записи регистра флагов после выполнения операции."],
        ],
        algorithm_rows=[
            ["1", "Блок alu_control сравнивает код операции с OP_OR, OP_NOR, OP_SRA и OP_INCS.", "Формируется внутренний код alu_op и сигнал записи флагов."],
            ["2", "Блоки OR, NOR, SRA и INCS параллельно вычисляют возможные результаты.", "Задержка выбора результата определяется мультиплексором, а не последовательным выполнением операций."],
            ["3", "Мультиплексор выбирает один из результатов по alu_op.", "На выход y_o поступает значение выбранной операции."],
            ["4", "Блок alu_flags анализирует выбранный результат и служебные признаки операции.", "Формируются Z, S, C и O."],
            ["5", "Диагностические выходы передают код операции, использование B и разрешение записи флагов.", "В RTL Viewer и на временной диаграмме видно, какая ветвь АЛУ активна."],
        ],
        module_rows=[
            ["lab_variant_pkg.vhd", "Общие константы варианта: разрядность, коды операций, коды внутренних режимов АЛУ."],
            ["alu_control.vhd", "Дешифрация кода операции и формирование alu_op, use_b, write_flags."],
            ["alu_or.vhd", "Комбинационный блок поразрядной операции OR."],
            ["alu_nor.vhd", "Комбинационный блок поразрядной операции NOR."],
            ["alu_sra.vhd", "Арифметический сдвиг вправо с выдачей вытолкнутого бита в C."],
            ["alu_incs.vhd", "Условный инкремент по флагу S с расчетом переноса и переполнения."],
            ["alu_flags.vhd", "Формирователь признаков результата Z, S, C и O."],
            ["lab5_alu_top.vhd", "Структурное объединение всех блоков АЛУ."],
            ["lab5_alu_rtl_view_top.vhd", "Верхний уровень с понятными именами пинов для RTL Viewer."],
            ["tb_lab5_alu.vhd", "Тестовая модель с assert-проверками всех операций."],
        ],
        implementation=[
            "Внутри АЛУ применена параллельная структура: каждый функциональный блок получает входные операнды одновременно. Это соответствует аппаратной реализации, в которой логические элементы OR/NOR, схема сдвига и сумматор INCS существуют как независимые ветви.",
            "Сумматор INCS расширяет операнд до 17 бит, чтобы сохранить перенос за старший разряд. Переполнение определяется для случая прибавления единицы к положительному числу, когда старший бит результата становится единицей. Для SRA переполнение не формируется, а переносом считается младший бит исходного слова.",
            "Формирователь флагов отделен от вычислительных блоков. Благодаря этому правила установки Z и S остаются общими для всех операций, а C и O передаются из той ветви, где они имеют смысл.",
        ],
        test_rows=[
            ["OR", "A=00F0h, B=0F0Fh, opcode=03h", "Y=0FFFh, Z=0, S=0"],
            ["NOR", "A=0FFFh, B=0FFFh, opcode=04h", "Y=F000h, S=1"],
            ["SRA", "A=8001h, opcode=05h", "Y=C000h, C=1, знак сохранен"],
            ["INCS", "A=C000h, FR.S=1, opcode=06h", "Y=C001h"],
            ["INCS без инкремента", "A=1234h, FR.S=0", "Y=1234h"],
            ["Нулевой результат", "A=FFFFh, B=FFFFh, opcode=04h", "Y=0000h, Z=1"],
        ],
        resource_rows=[
            ["Combinational ALUTs", "90"],
            ["Dedicated registers", "0"],
            ["Pins", "66"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ],
        conclusion=[
            "В ходе лабораторной работы разработано арифметико-логическое устройство для операций OR, NOR, SRA и INCS. Реализация разбита на независимые VHDL-блоки вычисления, управления и формирования флагов, что упрощает анализ схемы и соответствует аппаратной структуре АЛУ.",
            "Моделирование подтвердило корректность всех операций индивидуального варианта: логические операции дают ожидаемые поразрядные результаты, арифметический сдвиг сохраняет знак и формирует перенос, а INCS изменяет операнд только при установленном флаге S. Проект успешно компилируется в Quartus и готов для использования в составе микро-ЭВМ.",
        ],
    )

    lab6 = LabReport(
        num="6",
        title="Стековое запоминающее устройство",
        topic="«Стековое запоминающее устройство»",
        folder="lab6_stack",
        file_name="LR6_STACK_variant.docx",
        qpf="lab6_stack.qpf",
        top="lab6_stack_rtl_view_top",
        scheme="lab6_stack_scheme.png",
        wave="lab6_waveform.png",
        toc_pages=[
            ("1 Цель работы........................................................................", "3"),
            ("2 Исходные данные к работе....................................................", "4"),
            ("3 Теоретические сведения.......................................................", "5"),
            ("4 Выполнение работы..............................................................", "6"),
            ("5 Вывод....................................................................................", "10"),
        ],
        variant_rows=[
            ["Объем стека", "7 слов"],
            ["Разрядность слова", "16 бит"],
            ["Организация", "LIFO"],
            ["Направление роста", "Вниз, от старших адресов к младшим"],
            ["Начальное значение SP", "7, стек пуст"],
            ["Дополнительный режим", "Занесение результата АЛУ в стек"],
        ],
        theory=[
            "Стековое запоминающее устройство предназначено для временного хранения данных по принципу LIFO: последним записанное слово считывается первым. Такая организация используется для сохранения промежуточных результатов, адресов возврата, параметров процедур и данных, которые должны быть восстановлены в обратном порядке.",
            "В индивидуальном варианте стек содержит семь 16-разрядных слов. Направление роста выбрано вниз. Это означает, что при записи нового слова указатель стека уменьшается, а при чтении увеличивается. При пустом стеке SP равен 7, при полном стеке SP равен 0.",
            "Команда PUSH сначала проверяет наличие свободной ячейки, затем записывает слово в ячейку с номером SP-1 и уменьшает SP. Команда POP проверяет, что стек не пуст, считывает mem[SP] и увеличивает SP. Такой порядок обеспечивает корректное поведение при глубине, отличной от степени двойки.",
            "Для защиты от некорректных обращений используются диагностические флаги overflow и underflow. Overflow формируется при попытке PUSH в полный стек, underflow - при попытке POP из пустого стека. Состояния empty и full выводятся отдельно, чтобы блок управления мог запрещать недопустимые операции.",
        ],
        theory_table=[
            ["LIFO", "Правило доступа, при котором последним занесенное значение извлекается первым."],
            ["SP", "Указатель вершины стека. В данной реализации хранит индекс текущей вершины или 7 для пустого стека."],
            ["PUSH", "Запись слова в стек с уменьшением SP."],
            ["POP", "Чтение вершины стека с увеличением SP."],
            ["Overflow/underflow", "Диагностические признаки записи в полный стек и чтения из пустого стека."],
        ],
        task_rows=common_task_rows,
        signal_rows=[
            ["clock_i", "in", "1", "Тактовый сигнал синхронной части стека."],
            ["reset_i", "in", "1", "Сброс: очищает память, устанавливает SP=7 и empty=1."],
            ["stack_command_idle_push_pop_i", "in", "2", "Код операции: 00 - idle, 01 - PUSH, 10 - POP, 11 - PUSH_ALU."],
            ["data_from_register_file_i", "in", "16", "Данные из регистра общего назначения для обычного PUSH."],
            ["alu_result_for_push_i", "in", "16", "Результат АЛУ, который может быть занесен в стек режимом PUSH_ALU."],
            ["data_popped_to_register_file_o", "out", "16", "Слово, считанное с вершины стека при POP."],
            ["stack_pointer_current_value_o", "out", "3", "Текущее значение SP от 0 до 7."],
            ["status_stack_empty_o", "out", "1", "Признак пустого стека."],
            ["status_stack_full_o", "out", "1", "Признак полного стека."],
            ["error_push_to_full_stack_o", "out", "1", "Импульс ошибки PUSH при полном стеке."],
            ["error_pop_from_empty_stack_o", "out", "1", "Импульс ошибки POP при пустом стеке."],
            ["ctrl_push_enable_o", "out", "1", "Диагностический сигнал внутреннего разрешения записи."],
            ["ctrl_pop_enable_o", "out", "1", "Диагностический сигнал внутреннего разрешения чтения."],
        ],
        algorithm_rows=[
            ["1", "После reset регистровый массив очищается, SP получает значение 7.", "Стек пуст, empty=1, full=0."],
            ["2", "stack_control дешифрирует команду PUSH_REG, POP_REG или PUSH_ALU.", "Формируются внутренние сигналы push, pop и выбор источника данных."],
            ["3", "При PUSH и SP>0 входное слово записывается в mem[SP-1].", "SP уменьшается на единицу, новое слово становится вершиной."],
            ["4", "При POP и SP<7 считывается mem[SP].", "Считанное слово поступает на выход, SP увеличивается."],
            ["5", "При PUSH в полный стек или POP из пустого стека память и SP не изменяются.", "Формируются overflow или underflow."],
        ],
        module_rows=[
            ["lab_variant_pkg.vhd", "Общие константы: глубина стека, разрядность слова, коды команд."],
            ["stack_control.vhd", "Дешифратор команд стека и выбор источника данных."],
            ["stack7x16.vhd", "Синхронный регистровый массив стека, SP и флаги состояния."],
            ["lab6_stack_top.vhd", "Структурный верхний уровень стека."],
            ["lab6_stack_rtl_view_top.vhd", "Верхний уровень с развернутыми именами пинов для RTL Viewer."],
            ["tb_lab6_stack.vhd", "Тестовая модель PUSH, POP, PUSH_ALU, overflow и underflow."],
        ],
        implementation=[
            "Память стека описана массивом из семи 16-разрядных регистров. Поскольку глубина равна 7, указатель SP имеет 3 разряда и принимает значения от 0 до 7. Значение 7 не является адресом ячейки памяти, а служит служебным признаком пустого стека.",
            "Блок stack_control является комбинационным дешифратором. Он исключает одновременную запись и чтение, так как для каждого кода команды активируется только один из сигналов push или pop. Для команды PUSH_ALU дополнительно выбирается вход alu_data_i вместо reg_data_i.",
            "Флаги overflow и underflow являются импульсными: они устанавливаются только в том такте, где произошла запрещенная операция. Это удобно для блока управления, так как ошибка может быть зафиксирована или обработана отдельной логикой.",
        ],
        test_rows=[
            ["Reset", "rst_i=1", "SP=7, empty=1, full=0"],
            ["PUSH", "Запись 1111h и 2222h", "SP изменяется 7 -> 6 -> 5"],
            ["POP", "Чтение после двух PUSH", "На выходе 2222h, SP увеличивается до 6"],
            ["PUSH_ALU", "alu_data_i=C001h, cmd=11", "В стек заносится C001h"],
            ["Overflow", "PUSH при SP=0", "SP не меняется, overflow=1"],
            ["Underflow", "POP при SP=7", "SP не меняется, underflow=1"],
        ],
        resource_rows=[
            ["Combinational ALUTs", "70"],
            ["Dedicated registers", "136"],
            ["Pins", "61"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ],
        conclusion=[
            "В ходе лабораторной работы разработано стековое запоминающее устройство глубиной 7 слов по 16 бит. Реализованы операции PUSH, POP и дополнительный режим занесения результата АЛУ, что соответствует индивидуальному варианту и последующей интеграции с процессорным трактом.",
            "Моделирование подтвердило правило LIFO, корректное изменение SP при росте стека вниз, а также защиту от переполнения и чтения из пустого стека. Проект синтезируется в Quartus без ошибок и содержит верхний уровень с понятными пинами для просмотра в RTL Viewer.",
        ],
    )

    lab7 = LabReport(
        num="7",
        title="Арбитраж шин",
        topic="«Арбитраж шин»",
        folder="lab7_bus_arbiter",
        file_name="LR7_BUS_ARBITER_variant.docx",
        qpf="lab7_bus_arbiter.qpf",
        top="lab7_bus_rtl_view_top",
        scheme="lab7_bus_scheme.png",
        wave="lab7_waveform.png",
        toc_pages=[
            ("1 Цель работы........................................................................", "3"),
            ("2 Исходные данные к работе....................................................", "4"),
            ("3 Теоретические сведения.......................................................", "5"),
            ("4 Выполнение работы..............................................................", "6"),
            ("5 Вывод....................................................................................", "10"),
        ],
        variant_rows=[
            ["Способ арбитража", "Централизованный"],
            ["Тип подключения запросов", "Параллельные независимые линии REQ"],
            ["Правило обслуживания", "Фиксированный временной квант"],
            ["Количество ведущих устройств", "4"],
            ["Разрядность шины данных", "16 бит"],
            ["Прием данных ведомым", "Синхронный, по признаку занятости шины"],
        ],
        theory=[
            "Арбитраж шин необходим, когда несколько ведущих устройств могут обращаться к одному общему ресурсу. Если разрешить одновременную передачу, на общей шине возникнет конфликт: разные источники будут формировать разные значения. Поэтому вводится арбитр, который в каждый момент времени выдает разрешение только одному ведущему.",
            "В централизованной схеме все линии запросов поступают в один блок арбитража, а арбитр формирует линии предоставления GNT. Параллельное подключение означает, что запросы всех ведущих видны одновременно, и арбитр может принять решение без последовательного опроса внешних устройств.",
            "В данном варианте используется фиксированный временной квант. Арбитр циклически перебирает слоты 0, 1, 2 и 3. Каждый слот длится заданное число тактов. Если в текущем слоте соответствующий ведущий выставил запрос, ему выдается grant; если запроса нет, шина остается свободной до следующего слота.",
            "Такой алгоритм прост для аппаратной реализации и легко проверяется: одновременно активным может быть только один grant, а каждый активный ведущий получает возможность передачи в своем временном слоте.",
        ],
        theory_table=[
            ["REQ", "Линия запроса шины от ведущего устройства."],
            ["GNT", "Линия разрешения передачи, выдаваемая арбитром одному ведущему."],
            ["Slot", "Номер текущего временного окна обслуживания."],
            ["Quantum", "Длительность одного слота в тактах."],
            ["BUSY", "Признак занятости общей шины, используемый ведомым устройством как строб приема."],
        ],
        task_rows=common_task_rows,
        signal_rows=[
            ["clock_i", "in", "1", "Тактовый сигнал арбитра, ведущих и ведомого устройства."],
            ["reset_i", "in", "1", "Сброс счетчиков, слота и признаков обслуживания."],
            ["master_bus_request_lines_o", "out", "4", "Линии запросов четырех ведущих устройств."],
            ["arbiter_bus_grant_lines_o", "out", "4", "Линии разрешения шины; код one-hot."],
            ["shared_data_bus_to_slave_o", "out", "16", "Общая шина данных после мультиплексора."],
            ["slave_latched_data_from_bus_o", "out", "16", "Данные, зафиксированные ведомым устройством."],
            ["slave_data_valid_strobe_o", "out", "1", "Признак приема данных ведомым."],
            ["current_time_quantum_slot_o", "out", "2", "Номер текущего временного слота арбитра."],
            ["master_was_serviced_debug_o", "out", "4", "Диагностические признаки того, что каждый ведущий получил шину."],
        ],
        algorithm_rows=[
            ["1", "После reset слот устанавливается в 0, счетчик кванта сбрасывается.", "Арбитр начинает обслуживание с ведущего 0."],
            ["2", "Каждый ведущий независимо формирует REQ в заданном интервале тактов.", "На вход арбитра поступает вектор запросов req_i[3..0]."],
            ["3", "Арбитр проверяет запрос текущего слота.", "Если req(slot)=1, формируется единственный grant(slot)=1."],
            ["4", "Мультиплексор выбирает данные ведущего, которому выдан grant.", "На общей шине появляется слово 1111h, 2222h, 3333h или 4444h."],
            ["5", "Ведомое устройство фиксирует шину при активном busy.", "slave_data_o хранит последнее принятое слово, valid показывает факт приема."],
            ["6", "После завершения кванта счетчик переходит к следующему слоту.", "Порядок обслуживания повторяется циклически 0 -> 1 -> 2 -> 3."],
        ],
        module_rows=[
            ["master_device.vhd", "Модель ведущего устройства с собственным REQ, словом данных и признаком обслуживания."],
            ["bus_arbiter_parallel_quantum.vhd", "Центральный арбитр с фиксированным временным квантом."],
            ["bus_mux_4.vhd", "Мультиплексор общей 16-разрядной шины данных."],
            ["slave_sync.vhd", "Синхронное ведомое устройство, принимающее данные по BUSY."],
            ["lab7_bus_top.vhd", "Структурное соединение ведущих, арбитра, мультиплексора и ведомого."],
            ["lab7_bus_rtl_view_top.vhd", "Верхний уровень с понятными именами пинов для RTL Viewer."],
            ["tb_lab7_bus.vhd", "Проверка one-hot grant, соответствия REQ/GNT и обслуживания всех ведущих."],
        ],
        implementation=[
            "Ведущие устройства в лабораторной модели имеют разные моменты появления запроса и разные постоянные слова данных. Это сделано для того, чтобы на временной диаграмме было видно, какой ведущий получил шину и какое слово передано ведомому.",
            "Арбитр содержит регистр slot_r и счетчик quantum_r. Комбинационная часть формирует grant_s только для текущего слота. Если в этом слоте запроса нет, grant остается нулевым; арбитр не переключается на другого ведущего внутри того же кванта, потому что по заданию используется фиксированный временной квант.",
            "Общая шина реализована мультиплексором по вектору grant. Поскольку grant имеет one-hot-форму, на выход может быть передано слово только одного ведущего. Ведомое устройство синхронно запоминает значение общей шины, когда busy равен единице.",
        ],
        test_rows=[
            ["One-hot", "Все такты моделирования", "Активен не более один бит GNT"],
            ["REQ/GNT", "grant формируется только в активном слоте", "GNT не появляется без соответствующего REQ"],
            ["Квант", "QUANTUM_CYCLES=2", "slot повторяется два такта и затем увеличивается"],
            ["Передача данных", "grant0..grant3", "На шине появляются 1111h, 2222h, 3333h, 4444h"],
            ["Обслуживание", "Длительное моделирование", "Каждый ведущий получает шину хотя бы один раз"],
        ],
        resource_rows=[
            ["Combinational ALUTs", "34"],
            ["Dedicated registers", "19"],
            ["Pins", "49"],
            ["Block memory bits", "0"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ],
        conclusion=[
            "В ходе лабораторной работы разработана схема централизованного параллельного арбитража общей шины. Реализован фиксированный временной квант, четыре ведущих устройства, мультиплексор общей шины и синхронное ведомое устройство.",
            "Моделирование подтвердило отсутствие конфликтов на шине: в каждый момент времени активен не более один grant, данные на общей шине соответствуют выбранному ведущему, а каждый ведущий получает обслуживание в своем временном слоте. Проект синтезируется в Quartus и пригоден для дальнейшей интеграции с подсистемой памяти микро-ЭВМ.",
        ],
    )

    lab8 = LabReport(
        num="8",
        title="Организация кэш-памяти",
        topic="«Организация кэш-памяти»",
        folder="lab8_cache",
        file_name="LR8_CACHE_variant.docx",
        qpf="lab8_cache.qpf",
        top="lab8_cache_rtl_view_top",
        scheme="lab8_cache_scheme.png",
        wave="lab8_waveform.png",
        toc_pages=[
            ("1 Цель работы........................................................................", "3"),
            ("2 Исходные данные к работе....................................................", "4"),
            ("3 Теоретические сведения.......................................................", "5"),
            ("4 Выполнение работы..............................................................", "6"),
            ("5 Вывод....................................................................................", "10"),
        ],
        variant_rows=[
            ["Тип отображения", "k-мерное частично-ассоциативное"],
            ["Количество путей k", "4 строки в наборе"],
            ["Количество наборов", "16"],
            ["Размер строки", "1 слово, 16 бит"],
            ["Разбиение адреса", "TAG=A15..A4, SET=A3..A0"],
            ["Алгоритм замещения", "По наибольшей давности хранения"],
            ["Политика записи", "Write-through с обновлением строки кэша при попадании"],
        ],
        theory=[
            "Кэш-память используется для ускорения обращений процессора к основной памяти. В кэше хранятся копии недавно использованных слов, поэтому при повторном обращении процессор может получить данные без полного цикла доступа к ОЗУ.",
            "В k-мерном частично-ассоциативном кэше адрес выбирает набор, а внутри набора параллельно проверяются несколько строк. В данном варианте k равно 4, то есть каждый набор содержит четыре возможных места хранения слова. Всего используется 16 наборов.",
            "Адрес делится на поле TAG и поле SET. Младшие четыре бита A3..A0 выбирают один из 16 наборов. Старшие двенадцать бит A15..A4 хранятся как тег строки. Попадание возникает, если в выбранном наборе существует valid-строка с совпадающим тегом.",
            "Для замещения используется давность хранения. Каждая valid-строка имеет двухразрядный счетчик возраста. При загрузке новой строки ее возраст сбрасывается в ноль, а возраст остальных valid-строк выбранного набора увеличивается с насыщением до трех. Если свободной строки нет, вытесняется строка с максимальным возрастом.",
            "Политика записи выбрана write-through. Это означает, что каждое записываемое слово отправляется в основную память. Если запись попала в уже существующую строку кэша, строка обновляется; при промахе запись также может загрузить новую строку, чтобы последующее чтение этого адреса было попаданием.",
        ],
        theory_table=[
            ["TAG", "Старшая часть адреса, по которой проверяется принадлежность строки нужному адресу."],
            ["SET", "Младшие четыре бита адреса, выбирающие один из 16 наборов."],
            ["Valid", "Признак того, что строка содержит корректные данные."],
            ["Hit", "Совпадение тега при valid=1 в выбранном наборе."],
            ["Miss", "Отсутствие подходящей строки, требующее обращения к основной памяти."],
            ["Age", "Двухразрядный счетчик давности хранения строки для выбора жертвы."],
            ["Write-through", "Политика записи, при которой данные сразу передаются в основную память."],
        ],
        task_rows=common_task_rows,
        signal_rows=[
            ["clock_i", "in", "1", "Тактовый сигнал контроллера кэша и модели памяти."],
            ["reset_i", "in", "1", "Сброс valid-битов, тегов, данных, счетчиков возраста и автомата."],
            ["cpu_cache_request_i", "in", "1", "Запрос процессора на чтение или запись."],
            ["cpu_write_enable_i", "in", "1", "Тип операции: 0 - чтение, 1 - запись."],
            ["cpu_address_tag_set_i", "in", "16", "Адрес процессора; A15..A4 - тег, A3..A0 - номер набора."],
            ["cpu_write_data_to_cache_i", "in", "16", "Данные процессора для операции записи."],
            ["cpu_read_data_from_cache_o", "out", "16", "Данные, возвращаемые процессору при чтении."],
            ["cpu_cache_response_ready_o", "out", "1", "Признак готовности ответа кэша."],
            ["cache_hit_signal_o", "out", "1", "Импульс попадания."],
            ["cache_miss_signal_o", "out", "1", "Импульс промаха."],
            ["memory_request_from_cache_debug_o", "out", "1", "Диагностический запрос к основной памяти."],
            ["memory_write_enable_debug_o", "out", "1", "Диагностический признак записи в основную память."],
            ["memory_address_debug_o", "out", "16", "Адрес, передаваемый из кэша в основную память."],
            ["memory_data_debug_o", "out", "16", "Данные на интерфейсе основной памяти."],
        ],
        algorithm_rows=[
            ["1", "В состоянии IDLE при cpu_req_i=1 адрес разделяется на set и tag.", "Выбирается набор и запускается параллельная проверка четырех путей."],
            ["2", "Для каждого way проверяется valid_r(set,way)=1 и tag_r(set,way)=tag.", "При совпадении фиксируется hit_way."],
            ["3", "Если чтение попало в кэш, данные сразу выдаются процессору.", "cpu_ready=1, hit=1, обращение к RAM не требуется."],
            ["4", "Если чтение дало промах, выбирается свободная строка или строка с максимальным возрастом.", "Формируется запрос чтения RAM, автомат переходит к WAIT_READ_GRANT."],
            ["5", "После получения данных из RAM строка кэша заполняется.", "Записываются data, tag, valid, возраст новой строки сбрасывается."],
            ["6", "При записи данные всегда передаются в RAM.", "Реализуется write-through; при hit обновляется также строка кэша."],
            ["7", "После загрузки или записи новой строки обновляются счетчики возраста.", "Новая строка получает age=0, остальные valid-строки набора стареют до максимума 3."],
        ],
        module_rows=[
            ["lab_variant_pkg.vhd", "Общие параметры: 16 наборов, 4 пути, разрядность адреса и данных."],
            ["cache4way_age.vhd", "Контроллер кэша, массивы data/tag/valid/age, автомат обращения к RAM."],
            ["main_memory_sync.vhd", "Синхронная модель основной памяти для проверки hit/miss/write-through."],
            ["lab8_cache_top.vhd", "Структурное соединение кэша и основной памяти."],
            ["lab8_cache_rtl_view_top.vhd", "Верхний уровень с понятными именами пинов для RTL Viewer."],
            ["tb_lab8_cache.vhd", "Тестовая модель чтения, записи, попадания, промаха и замещения."],
        ],
        implementation=[
            "В кэше используются четыре двумерных массива: data_r хранит слова данных, tag_r хранит теги, valid_r показывает корректность строки, age_r хранит давность. Индексация массивов выполняется по номеру набора и номеру пути.",
            "Контроллер реализован конечным автоматом с состояниями IDLE, WAIT_READ_GRANT, WAIT_READ_DATA и WAIT_WRITE_GRANT. В состоянии IDLE выполняется проверка попадания и выбор жертвы. При промахе чтения автомат запрашивает основную память и ожидает данные. При записи автомат формирует write-through-запрос.",
            "Модель основной памяти main_memory_sync возвращает grant синхронно с запросом и хранит 256 слов. Начальное содержимое выбрано детерминированным, поэтому ожидаемые значения на временной диаграмме легко проверить: адрес 0010h возвращает 1110h, 0020h возвращает 1220h и так далее.",
        ],
        test_rows=[
            ["Read miss", "Первое чтение R[0010h]", "Промах, загрузка 1110h из RAM"],
            ["Read hit", "Повторное чтение R[0010h]", "Попадание, данные возвращаются из кэша"],
            ["Write-through", "Запись W[0010h]=ABCDh", "RAM обновлена, строка кэша при hit обновлена"],
            ["Read after write", "Чтение R[0010h] после записи", "Попадание, возвращается ABCDh"],
            ["Заполнение набора", "R[0020h], R[0030h], R[0040h], R[0050h]", "Один набор заполняется четырьмя путями"],
            ["Replacement", "Повтор R[0010h] после вытеснения", "Промах, затем возврат ABCDh из RAM"],
        ],
        resource_rows=[
            ["Combinational ALUTs", "1351"],
            ["Dedicated registers", "2120"],
            ["Pins", "89"],
            ["Block memory bits", "4096"],
            ["Результат компиляции", "Full compilation: 0 errors"],
        ],
        conclusion=[
            "В ходе лабораторной работы разработана 4-way set associative кэш-память с 16 наборами, разбиением адреса на TAG и SET, valid-битами, тегами, данными и счетчиками давности хранения.",
            "Моделирование подтвердило чтение при попадании, обработку промаха с загрузкой из основной памяти, сквозную запись write-through и замещение строки по наибольшей давности хранения. Проект успешно синтезируется в Quartus и может использоваться как подсистема памяти в составе микро-ЭВМ.",
        ],
    )

    return [lab5, lab6, lab7, lab8]


def make_report(lab: LabReport) -> Path:
    lab = prepare_variant4_lab(lab)
    doc = Document()
    configure_doc(doc)
    title_page(doc, lab)
    add_common_opening(doc, lab)
    add_theory(doc, lab)
    add_execution(doc, lab)
    add_conclusion(doc, lab)

    out = ROOT / lab.folder / "report" / lab.file_name
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out)
    except PermissionError:
        out = out.with_name(f"{out.stem}_final{out.suffix}")
        doc.save(out)
    return out


def main():
    from build_variant4_assets import main as build_variant4_assets

    build_variant4_assets()
    generated: list[Path] = []
    for lab in report_data():
        generated.append(make_report(lab))
    print("Generated reports:")
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
