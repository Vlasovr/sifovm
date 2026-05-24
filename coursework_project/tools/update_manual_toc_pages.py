from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "documents" / "Пояснительная_записка.docx"
PDF = ROOT / "documents" / "Пояснительная_записка.pdf"
PDFTOTEXT = Path(
    r"C:\Users\user\AppData\Local\Microsoft\WinGet\Packages"
    r"\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
    r"\poppler-25.07.0\Library\bin\pdftotext.exe"
)
PDFINFO = Path(
    r"C:\Users\user\AppData\Local\Microsoft\WinGet\Packages"
    r"\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
    r"\poppler-25.07.0\Library\bin\pdfinfo.exe"
)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def page_count(pdf: Path) -> int:
    out = subprocess.check_output([str(PDFINFO), str(pdf)], text=True, encoding="utf-8", errors="ignore")
    m = re.search(r"Pages:\s+(\d+)", out)
    if not m:
        raise RuntimeError("Cannot read PDF page count")
    return int(m.group(1))


def extract_pages(pdf: Path) -> dict[int, str]:
    pages = {}
    for page in range(1, page_count(pdf) + 1):
        text = subprocess.check_output(
            [str(PDFTOTEXT), "-layout", "-f", str(page), "-l", str(page), str(pdf), "-"],
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
        pages[page] = norm(text)
    return pages


def search_key(title: str) -> str:
    title = norm(title)
    if title.startswith("ПРИЛОЖЕНИЕ "):
        m = re.match(r"(ПРИЛОЖЕНИЕ\s+[А-ЯЁ])", title)
        return m.group(1) if m else title
    if title.startswith("СПИСОК "):
        return "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
    if title.startswith("ЗАКЛЮЧЕНИЕ"):
        return "ЗАКЛЮЧЕНИЕ"
    m = re.match(r"(\d+(?:\.\d+)*)(?:\s+)(.+)", title)
    if m:
        number = m.group(1)
        if number == "3.7":
            return "3.7 Функциональное моделирование арифметико-логического устройства"
        if number == "3.10":
            return "3.10 Функциональное моделирование контроллера прямого доступа к памяти"
        return title
    return title


def find_page(title: str, pages: dict[int, str]) -> int | None:
    key = search_key(title)
    key_norm = norm(key)
    # Skip the title page and manual TOC pages.
    for page, text in pages.items():
        if page <= 4:
            continue
        if key_norm in text:
            return page

    # Fallback for long wrapped headings: use the section number and first words.
    m = re.match(r"(\d+(?:\.\d+)*)\s+(.+)", key_norm)
    if m:
        num = m.group(1)
        words = m.group(2).split()[:3]
        for page, text in pages.items():
            if page <= 4:
                continue
            if num in text and all(word in text for word in words):
                return page
    return None


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def main() -> None:
    document = Document(DOCX)
    pages = extract_pages(PDF)
    updates = []

    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if "\t" not in text:
            continue
        title, old_page = text.rsplit("\t", 1)
        if not old_page.strip().isdigit():
            continue
        page = find_page(title, pages)
        if page is None:
            continue
        new_text = f"{title}\t{page}"
        if new_text != text:
            updates.append((idx, text, new_text))
            set_paragraph_text(paragraph, new_text)

    document.save(DOCX)
    print(f"Updated {len(updates)} TOC lines")
    for idx, old, new in updates:
        print(f"{idx}: {old} -> {new}")


if __name__ == "__main__":
    main()
