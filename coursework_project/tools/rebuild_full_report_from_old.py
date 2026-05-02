# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
OLD_REPORT = ROOT / "_archive_unused_2026-05-02" / "Курсовая" / "Записка Власов.docx"
DOC_DIR = ROOT / "coursework_project" / "documents"
OUT_REPORT = DOC_DIR / "Пояснительная_записка.docx"
FULL_COPY = DOC_DIR / "Пояснительная_записка_полная.docx"
SHORT_BACKUP = DOC_DIR / "Пояснительная_записка_краткая_архив.docx"


REPLACEMENTS = [
    ("в соответствии с вариантом 4 технического задания", "в соответствии с индивидуальным техническим заданием"),
    ("в соответствии с вариантом 4", "в соответствии с индивидуальным техническим заданием"),
    ("вариантом 4 технического задания", "индивидуальным техническим заданием"),
    ("В рамках выбранного варианта", "В рамках индивидуального задания"),
    ("Согласно варианту задания", "Согласно техническому заданию"),
    ("варианта 4", "технического задания"),
    ("вариант 4", "индивидуальное техническое задание"),
    ("данном варианте", "данном проекте"),
    ("в данном варианте", "в данном проекте"),
    ("выбранного варианта", "индивидуального задания"),
    ("параметров варианта", "параметров задания"),
    ("таблице варианта", "техническому заданию"),
    ("по условию варианта", "по условию задания"),
    ("заданным вариантом 4", "заданным техническим заданием"),
    ("заданный вариантом 4", "заданный техническим заданием"),
    ("заданы схем", "заданы схем"),
    ("заданы схема", "задана схема"),
    ("Тип адресации, заданный техническим заданием", "Тип адресации, заданный техническим заданием"),
    ("обязательные по варианту задания", "обязательные по техническому заданию"),
    ("операции варианта INCS, OR, NOR и SRA", "заданные операции INCS, OR, NOR и SRA"),
    ("операции варианта", "заданные операции"),
    ("Проявление в индивидуальное техническое задание", "Проявление в проекте"),
    ("Проявление в проекте 4", "Проявление в проекте"),
    ("Проявление в варианте 4", "Проявление в проекте"),
    ("микропрограммного варианта устройства управления", "микропрограммного устройства управления"),
    ("микропрограммный вариант УУ", "микропрограммное УУ"),
    ("микропрограммного варианта УУ", "микропрограммного УУ"),
    ("в таком варианте", "при таком подходе"),
    ("базовом варианте", "базовой реализации"),
    ("текущего компактного варианта 4", "текущей компактной реализации"),
    ("Содержание и варианты заданий", "Содержание заданий"),
    ("варианта", "задания"),
    ("варианте", "проекте"),
    ("вариантом", "заданием"),
    ("вариант", "задание"),
    ("Variant", "Project"),
    ("variant", "project"),
]


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def clear_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        p_pr.remove(p_pr.numPr)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    new_para.add_run(text)
    return new_para


def replace_text_in_paragraph(paragraph) -> None:
    original = paragraph.text
    if not original:
        return
    updated = original
    for old, new in REPLACEMENTS:
        updated = updated.replace(old, new)
    updated = re.sub(r"\s+4(?=[,.;:])", "", updated)
    if updated != original:
        set_paragraph_text(paragraph, updated)


def replace_everywhere(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.start_type = WD_SECTION_START.NEW_PAGE


def normalize_styles(doc: Document) -> None:
    for style_name in ("Normal", "Body Text", "List Paragraph"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(14)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.bold = True
            style.font.size = Pt(14)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(3)
        if paragraph.style.name == "Normal" and paragraph.text.strip():
            paragraph.paragraph_format.first_line_indent = Cm(1.25)


def remove_leading_and_trailing_empty_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            break
        delete_paragraph(paragraph)

    for paragraph in reversed(list(doc.paragraphs)):
        if paragraph.text.strip():
            break
        delete_paragraph(paragraph)


def prepend_front_matter(doc: Document) -> None:
    first = doc.paragraphs[0]
    items = [
        ("Министерство образования Республики Беларусь", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("Учреждение образования", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("Белорусский государственный университет информатики и радиоэлектроники", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("Кафедра электронных вычислительных машин", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER),
        ("к курсовому проекту на тему", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("«Разработка микро-ЭВМ на ПЛИС»", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("Студент: Власов", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("Руководитель: ____________________", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("Минск 2026", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("PAGE_BREAK", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
        ("СОДЕРЖАНИЕ", "Heading 1", WD_ALIGN_PARAGRAPH.CENTER),
        ("Введение", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("1 Разработка общей структуры микро-ЭВМ", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("2 Разработка основных устройств микро-ЭВМ", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("3 Функциональное моделирование", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("4 Анализ и оптимизация разработанной микро-ЭВМ", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("Заключение", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("Список использованных источников", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("Приложения", "Normal", WD_ALIGN_PARAGRAPH.LEFT),
        ("PAGE_BREAK", "Normal", WD_ALIGN_PARAGRAPH.CENTER),
    ]

    for text, style, align in items:
        para = first.insert_paragraph_before("" if text == "PAGE_BREAK" else text, style=style)
        para.alignment = align
        if text == "PAGE_BREAK":
            para.add_run().add_break(WD_BREAK.PAGE)

    for para in doc.paragraphs:
        if para.text.strip() in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"}:
            try:
                para.style = "Heading 1"
            except KeyError:
                pass
            clear_paragraph_numbering(para)


def patch_headings_and_alu(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "РАЗРАБОТКА ОБЩЕЙ СТРУКТУРЫ МИКРО-ЭВМ":
            set_paragraph_text(paragraph, "1 РАЗРАБОТКА ОБЩЕЙ СТРУКТУРЫ МИКРО-ЭВМ")
            try:
                paragraph.style = "Heading 1"
            except KeyError:
                pass
            clear_paragraph_numbering(paragraph)
        elif text == "Функциональный состав микро-ЭВМ.":
            set_paragraph_text(paragraph, "1.1 Функциональный состав микро-ЭВМ")
            try:
                paragraph.style = "Heading 2"
            except KeyError:
                pass
            clear_paragraph_numbering(paragraph)
        elif text == "Арифметико-логическое устройство":
            set_paragraph_text(paragraph, "2.3 Арифметико-логическое устройство и блок регистров общего назначения")
            try:
                paragraph.style = "Heading 2"
            except KeyError:
                pass
            clear_paragraph_numbering(paragraph)

    paras = list(doc.paragraphs)
    start = None
    end = None
    for i, paragraph in enumerate(paras):
        text = paragraph.text.strip()
        if text.startswith("Данный блок выполняет арифметические"):
            start = i
        if text.startswith("ALU_R2[7..0]"):
            end = i
            break

    if start is not None and end is not None and end >= start:
        replacement = [
            "Арифметико-логическое устройство выполняет операции INCS, OR, NOR и SRA над 16-разрядными словами. Операнды поступают из блока регистров общего назначения или из тракта данных памяти, а результат возвращается в выбранный регистр через внутреннюю шину процессора.",
            "Входные сигналы АЛУ: A[15:0] и B[15:0] — операнды; SFLAG — текущее значение признака знака для операции INCS; ALU_OP[2:0] — код выполняемой операции. Выходные сигналы: Y[15:0] — результат; Z, S, C, O — признаки нуля, знака, переноса/вытесненного бита и переполнения.",
            "Для операции SRA сохраняется старший бит исходного операнда, а младший бит заносится во флаг C. Для OR и NOR используются оба операнда, при этом C и O сбрасываются. Для INCS к операнду прибавляется текущее значение флага S, поэтому при S = 0 регистр не изменяется, а при S = 1 увеличивается на единицу.",
            "Электрическая схема арифметико-логического устройства приведена в графической части проекта.",
            "Рис. 2.3 - Арифметико-логическое устройство.",
        ]
        set_paragraph_text(paras[start], replacement[0])
        last = paras[start]
        for text in replacement[1:]:
            last = insert_after(last, text, "Normal")
        for paragraph in paras[start + 1 : end + 1]:
            if paragraph.text.strip() != "Блок регистров общего назначения":
                delete_paragraph(paragraph)


def patch_intro(doc: Document) -> None:
    paras = list(doc.paragraphs)
    start = None
    end = None
    for i, paragraph in enumerate(paras):
        text = paragraph.text.strip()
        if text == "ВВЕДЕНИЕ":
            start = i
            continue
        if start is not None and text.startswith("1 РАЗРАБОТКА"):
            end = i
            break
    if start is None or end is None:
        return

    intro_items = [
        ("Целью курсового проекта является разработка микро-ЭВМ на ПЛИС в соответствии с индивидуальным техническим заданием по дисциплине «Структурная и функциональная организация ЭВМ». Введение фиксирует основные исходные параметры системы, от которых далее зависят структура команд, состав функциональных блоков, набор управляющих сигналов и перечень временных диаграмм.", "Normal"),
        ("Разрабатываемая микро-ЭВМ имеет следующие основные магистрали:", "Normal"),
        ("- шина данных разрядностью 16 бит;", "Normal"),
        ("- адресная шина разрядностью 16 бит;", "Normal"),
        ("- управляющая шина, разрядность которой определяется набором сигналов чтения и записи памяти, разрешения регистров, выбора источников внутренней шины, управления АЛУ, стеком, кэшем, КПДП, арбитром и предсказателем переходов.", "Normal"),
        ("Архитектура микро-ЭВМ принята гарвардской: ПЗУ команд и ОЗУ данных имеют раздельные тракты обращения. Это позволяет выборке команды не конфликтовать с обычным доступом к данным, однако обращения процессора и КПДП к ОЗУ всё равно требуют арбитража. Команды имеют фиксированный двухсловный формат: первое слово содержит код операции и номер регистра, второе слово используется как прямой адрес операнда, адрес перехода или резервное поле.", "Normal"),
        ("Тип адресации в проекте прямо-регистровый. Операнд или приёмник результата задаётся номером одного из 12 регистров общего назначения R0-R11, а адрес ячейки памяти располагается во втором слове команды. Для команд JMP и JZ второе слово содержит адрес следующей команды в пространстве ПЗУ.", "Normal"),
        ("Арифметико-логическое устройство должно выполнять операции INCS, OR, NOR и SRA. Операция INCS использует значение флага S, операции OR и NOR выполняют побитовую логическую обработку, а SRA реализует арифметический сдвиг вправо с сохранением знакового бита. Результаты операций сопровождаются формированием флагов Z, S, C и O.", "Normal"),
        ("Кэш-память данных организуется как четырёхканальная множественно-ассоциативная структура с вытеснением строки по наибольшей давности хранения. Синхронизация с основной памятью выполняется по схеме сквозной записи: при записи слово согласуется и с кэшем, и с ОЗУ, поэтому последующее чтение получает актуальное значение.", "Normal"),
        ("Арбитр шин реализуется как централизованный параллельный блок. Его назначение - исключить одновременное управление общей магистралью данных со стороны процессорной части и контроллера прямого доступа к памяти. Блок предсказания переходов выполняется отдельно от устройства управления и использует схему A4 с индексированием по сочетанию младших разрядов счётчика команд и регистра глобальной истории GHR.", "Normal"),
        ("Стек реализуется отдельным блоком глубиной 7 слов по 16 бит с направлением роста вниз. Контроллер прямого доступа к памяти передаёт в ОЗУ блок объёмом 6 байт, что при словной организации памяти соответствует трём 16-разрядным словам, начиная с адреса 000Ah.", "Normal"),
        ("Разрабатываемая микро-ЭВМ является многотактовой, а не конвейерной. Поэтому выполнение команды рассматривается как последовательность фаз выборки, декодирования, обращения к операндам, исполнения и записи результата. Такой подход упрощает проверку временных диаграмм и позволяет отдельно показать работу ПЗУ, ОЗУ, АЛУ, стека, кэша, КПДП, арбитра и предсказателя переходов.", "Normal"),
        ("Разработка и проверка проекта выполняются в среде Altera Quartus II 9.1 с использованием VHDL-описаний и функционального моделирования. Итогом работы должна стать согласованная структура микро-ЭВМ, набор приложений с графическими схемами и листингами, а также результаты тестирования, подтверждающие корректность работы отдельных блоков и всей системы.", "Normal"),
    ]

    heading = paras[start]
    for paragraph in paras[start + 1 : end]:
        delete_paragraph(paragraph)

    last = heading
    for text, style in intro_items:
        last = insert_after(last, text, style)


def block_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def find_body_index(doc: Document, needle: str) -> int | None:
    for i, element in enumerate(doc.element.body):
        if element.tag == qn("w:p") and block_text(element) == needle:
            return i
    return None


def add_body_paragraph(doc: Document, text: str, style: str = "Normal", indent: bool = True):
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(3)
    if style == "Normal" and indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    return paragraph


def add_list_line(doc: Document, text: str, level: int = 0):
    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.7 + 0.55 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(1)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        run.bold = True
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def build_chapter_1_1_1_2() -> Document:
    temp = Document()
    add_body_paragraph(temp, "1.1 Функциональный состав микро-ЭВМ", "Heading 2", indent=False)

    p = temp.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    r = p.add_run("Корневыми блоками микро-ЭВМ являются:")
    r.bold = True
    r.underline = True

    add_list_line(temp, "- блок запоминающих устройств:")
    add_list_line(temp, "- ПЗУ программ для хранения двухсловных команд;", 1)
    add_list_line(temp, "- ОЗУ данных для хранения операндов, результатов и области обмена с КПДП;", 1)
    add_list_line(temp, "- кэш-память данных со схемой управления, тегами, признаками действительности и счётчиками давности;", 1)
    add_list_line(temp, "- схема согласования обращений к памяти через арбитр шин.", 1)

    add_list_line(temp, "- блок устройства управления:")
    add_list_line(temp, "- регистр текущей команды IR0/IR1;", 1)
    add_list_line(temp, "- декодер кода операции и номера регистра;", 1)
    add_list_line(temp, "- фазовый автомат выполнения команды;", 1)
    add_list_line(temp, "- схема выбора следующего значения указателя команд IP;", 1)
    add_list_line(temp, "- формирователь управляющих сигналов для РОН, АЛУ, стека, кэша, ОЗУ, КПДП и предсказателя переходов.", 1)

    add_list_line(temp, "- блок регистров общего назначения:")
    add_list_line(temp, "- 12 регистров R0-R11 разрядностью 16 бит;", 1)
    add_list_line(temp, "- два порта чтения и один порт записи для выдачи операндов и фиксации результата.", 1)

    add_list_line(temp, "- блок арифметико-логического устройства:")
    add_list_line(temp, "- узлы операций OR и NOR;", 1)
    add_list_line(temp, "- узел арифметического сдвига вправо SRA;", 1)
    add_list_line(temp, "- узел инкремента по флагу S для операции INCS;", 1)
    add_list_line(temp, "- схема формирования флагов Z, S, C и O.", 1)

    add_list_line(temp, "- блок стека:")
    add_list_line(temp, "- стековая память глубиной 7 слов;", 1)
    add_list_line(temp, "- указатель вершины SP с направлением роста вниз;", 1)
    add_list_line(temp, "- схемы контроля пустого и полного состояния.", 1)

    add_list_line(temp, "- блок общих операций и служебных подсистем:")
    add_list_line(temp, "- централизованный параллельный арбитр общей шины данных;", 1)
    add_list_line(temp, "- контроллер прямого доступа к памяти;", 1)
    add_list_line(temp, "- блок предсказания переходов по схеме A4;", 1)
    add_list_line(temp, "- мультиплексоры адресов и данных для выбора активного источника обмена.", 1)

    add_body_paragraph(
        temp,
        "Арбитр шин, предсказатель переходов и контроллер прямого доступа к памяти рассматриваются как отдельные функциональные блоки. Они не заменяют устройство управления, а дополняют его: арбитр разрешает конфликты доступа к ОЗУ, предсказатель выдаёт предполагаемый адрес следующей команды, а КПДП выполняет обмен с памятью без участия арифметико-логического тракта процессора.",
    )

    p = temp.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    r = p.add_run("Микро-ЭВМ помимо регистров общего назначения располагает специальными регистрами:")
    r.bold = True
    r.underline = True

    add_list_line(temp, "- IP (instruction pointer) - 16-разрядный указатель команды, задающий адрес первого слова текущей или следующей команды в ПЗУ;")
    add_list_line(temp, "- IR0, IR1 (instruction register) - регистры двух слов команды; IR0 хранит код операции и номер регистра, IR1 хранит адресную часть;")
    add_list_line(temp, "- FR (flag register) - регистр флагов Z, S, C и O, отражающий состояние результата последней арифметико-логической операции;")
    add_list_line(temp, "- SP (stack pointer) - указатель вершины стека, изменяемый при выполнении PUSH и POP;")
    add_list_line(temp, "- GHR (global history register) - регистр глобальной истории переходов, используемый предсказателем;")
    add_list_line(temp, "- служебные регистры КПДП - текущий адрес передачи, счётчик оставшихся байт и буфер принимаемого слова;")
    add_list_line(temp, "- служебные регистры кэша - теги, признаки действительности строк и счётчики давности хранения.")

    add_body_paragraph(
        temp,
        "Специальные регистры не адресуются командами как обычные регистры R0-R11. Их содержимое задаётся аппаратной логикой соответствующих блоков и используется для хранения промежуточного состояния вычислительного процесса. Для хранения пользовательских данных должны использоваться РОН или ОЗУ данных, так как достоверность содержимого специальных регистров определяется текущей фазой выполнения команды.",
    )

    add_body_paragraph(temp, "1.2 Разработка системы команд", "Heading 2", indent=False)

    p = temp.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    r = p.add_run("Количество реализованных команд - 11:")
    r.bold = True
    r.underline = True

    for line in [
        "- HLT - останов выполнения программы;",
        "- MOV adr, Rn - чтение слова из памяти данных в регистр общего назначения;",
        "- MOV Rn, adr - запись слова из регистра общего назначения в память данных;",
        "- OR Rn, adr - побитовое ИЛИ регистра и слова памяти;",
        "- NOR Rn, adr - инверсия результата побитового ИЛИ;",
        "- SRA Rn - арифметический сдвиг регистра вправо;",
        "- INCS Rn - инкремент регистра на значение флага S;",
        "- PUSH Rn - занесение регистра в стек;",
        "- POP Rn - извлечение слова из стека в регистр;",
        "- JMP adr - безусловная загрузка адреса в IP;",
        "- JZ adr - условный переход при установленном флаге Z.",
    ]:
        add_list_line(temp, line)

    add_body_paragraph(
        temp,
        "Для кодирования одиннадцати команд достаточно четырёх бит, однако в проекте поле кода операции принято восьмибитным. Такое решение согласуется с 16-разрядной организацией слова команды, упрощает декодирование в VHDL и оставляет резерв кодов для возможного расширения системы команд. Неиспользуемые коды рассматриваются как зарезервированные и в аппаратной реализации переводят процессор в безопасное состояние останова.",
    )
    add_body_paragraph(
        temp,
        "Команда имеет фиксированную длину два 16-разрядных слова. Первое слово содержит код операции и номер регистра, второе слово используется как адрес операнда в ОЗУ, адрес перехода в ПЗУ или резервное поле для команд, которым адрес не требуется. Фиксированный формат выбран для того, чтобы фазовый автомат всегда выполнял одинаковую процедуру выборки IR0 и IR1.",
    )

    p = temp.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    r = p.add_run("Положение полей команды фиксировано:")
    r.bold = True
    r.underline = True

    add_list_line(temp, "- OPCODE[7..0] - код операции, всегда размещается в разрядах IR0(15..8);")
    add_list_line(temp, "- REG_NUM[3..0] - номер регистра общего назначения, размещается в разрядах IR0(7..4);")
    add_list_line(temp, "- X[3..0] - служебное поле IR0(3..0), не влияющее на выполнение текущего набора команд;")
    add_list_line(temp, "- ADDR[15..0] - адресная часть IR1(15..0), используемая как адрес ОЗУ, адрес перехода в ПЗУ или резервное поле;")
    add_list_line(temp, "- RESERVED - коды и поля, оставленные для дальнейшего расширения системы команд.")

    add_body_paragraph(
        temp,
        "Поле REG_NUM имеет разрядность 4 бита, поэтому теоретически позволяет адресовать шестнадцать регистров. В проектируемой микро-ЭВМ фактически используются регистры R0-R11, а коды 12-15 не применяются в тестовой программе и считаются зарезервированными.",
    )
    add_body_paragraph(
        temp,
        "Адресное поле второго слова интерпретируется в зависимости от команды. Для MOV, OR и NOR оно задаёт адрес в ОЗУ данных; для JMP и JZ - адрес команды в ПЗУ; для HLT, SRA, INCS, PUSH и POP второе слово выбирается из ПЗУ вместе с командой, но при исполнении не используется.",
    )

    p = temp.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None
    r = p.add_run("Кодировка КОП принята следующей:")
    r.bold = True
    r.underline = True

    for line in [
        "- 00h - HLT;",
        "- 01h - MOV adr, Rn;",
        "- 02h - MOV Rn, adr;",
        "- 03h - OR Rn, adr;",
        "- 04h - NOR Rn, adr;",
        "- 05h - SRA Rn;",
        "- 06h - INCS Rn;",
        "- 07h - PUSH Rn;",
        "- 08h - POP Rn;",
        "- 09h - JMP adr;",
        "- 0Ah - JZ adr.",
    ]:
        add_list_line(temp, line)

    add_body_paragraph(
        temp,
        "Любой другой код операции считается зарезервированным. В аппаратной модели такой код не передаётся на выполнение произвольному блоку, а приводит процессор к безопасному останову, что упрощает отладку и защищает модель от случайного выполнения неопределённой команды.",
    )

    add_body_paragraph(
        temp,
        "Принятая система команд делится на четыре группы: команды обмена с памятью, арифметико-логические команды, стековые команды и команды передачи управления. Команда JZ включена для содержательной проверки предсказателя переходов: без условного перехода предсказатель невозможно показать в работе, так как безусловный JMP не требует анализа результата вычисления.",
    )
    return temp


def patch_chapter_1_1_1_2(doc: Document) -> None:
    body = doc.element.body
    start = find_body_index(doc, "1.1 Функциональный состав микро-ЭВМ")
    end = find_body_index(doc, "Описание взаимодействия всех блоков микро-ЭВМ при выполнении команд программы.")
    if start is None or end is None or end <= start:
        return

    for element in list(body)[start:end]:
        body.remove(element)

    temp = build_chapter_1_1_1_2()
    insert_before = list(body)[start]
    for element in list(temp.element.body):
        if element.tag == qn("w:sectPr"):
            continue
        body.insert(body.index(insert_before), deepcopy(element))


def patch_chapter_1_3_heading_cleanup(doc: Document) -> None:
    old_title = "Описание взаимодействия всех блоков микро-ЭВМ при выполнении команд программы."
    new_title = "1.3 Описание взаимодействия всех блоков микро-ЭВМ при выполнении команд программы"

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == old_title:
            set_paragraph_text(paragraph, new_title)
            try:
                paragraph.style = "Heading 2"
            except KeyError:
                pass
            clear_paragraph_numbering(paragraph)
            break

    body = doc.element.body
    start = find_body_index(doc, new_title)
    end = find_body_index(doc, "Общая последовательность выполнения команды")
    if start is None or end is None or end <= start:
        return

    for element in list(body)[start + 1 : end]:
        body.remove(element)


def append_application_list(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    h = doc.add_paragraph("ПРИЛОЖЕНИЯ")
    try:
        h.style = "Heading 1"
    except KeyError:
        pass
    rows = [
        ("Приложение А", "Графическая часть: структурные и функциональные схемы микро-ЭВМ."),
        ("Приложение Б", "Листинг тестовой программы, дампы памяти и контрольные данные."),
        ("Приложение В", "VHDL-листинги основных модулей проекта."),
        ("Приложение Г", "Материалы функционального моделирования и контрольные временные диаграммы."),
        ("Приложение Д", "Инструкция по снятию временных диаграмм в Quartus/ModelSim."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Обозначение"
    table.rows[0].cells[1].text = "Содержание"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def iter_all_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def finalize_formatting(doc: Document) -> None:
    blue = RGBColor(47, 117, 181)
    black = RGBColor(0, 0, 0)
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text.strip()
        if text.startswith("1.1 ") or text.startswith("1.2 ") or text.startswith("1.3 "):
            try:
                paragraph.style = "Heading 2"
            except KeyError:
                pass
        is_heading = paragraph.style.name.startswith("Heading")
        if is_heading:
            clear_paragraph_numbering(paragraph)
            paragraph.paragraph_format.first_line_indent = None
            if paragraph.style.name == "Heading 1":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(14)
            run.font.color.rgb = blue if is_heading else black
            if is_heading:
                run.font.bold = True


def clone_old_report() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    if OUT_REPORT.exists() and not SHORT_BACKUP.exists():
        shutil.copy2(OUT_REPORT, SHORT_BACKUP)
    shutil.copy2(OLD_REPORT, FULL_COPY)
    return FULL_COPY


def assert_clean(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    bad = re.search(r"variant|вариант", text, flags=re.IGNORECASE)
    if bad:
        raise RuntimeError(f"Forbidden leftover token found: {bad.group(0)!r}")


def main() -> None:
    if not OLD_REPORT.exists():
        raise FileNotFoundError(OLD_REPORT)

    working = clone_old_report()
    doc = Document(str(working))
    set_margins(doc)
    normalize_styles(doc)
    remove_leading_and_trailing_empty_paragraphs(doc)
    patch_headings_and_alu(doc)
    replace_everywhere(doc)
    patch_intro(doc)
    patch_chapter_1_1_1_2(doc)
    patch_chapter_1_3_heading_cleanup(doc)
    prepend_front_matter(doc)
    append_application_list(doc)
    finalize_formatting(doc)
    doc.save(str(working))
    assert_clean(working)

    try:
        shutil.copy2(working, OUT_REPORT)
    except PermissionError:
        print(f"Main report is open or locked, left full copy at: {working}")
    else:
        print(f"Updated: {OUT_REPORT}")
    print(f"Full copy: {working}")
    print(f"Short backup: {SHORT_BACKUP}")


if __name__ == "__main__":
    main()
