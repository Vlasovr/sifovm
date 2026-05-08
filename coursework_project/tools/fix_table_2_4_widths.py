from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT = Path("D:/git/sifovm/coursework_project/documents/Пояснительная_записка.docx")


def set_fixed_table(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = Cm(widths_cm[idx])
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_cm[idx] * 567)))
            tc_w.set(qn("w:type"), "dxa")

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def main() -> None:
    doc = Document(str(REPORT))
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "S_FETCH0_REQ" in text and "S_FETCH1_REQ" in text and "Основные состояния RTL" in text:
            set_fixed_table(table, [1.25, 7.15, 7.15])
            break
    else:
        raise RuntimeError("Table 2.4 not found")

    doc.save(str(REPORT))
    print(f"Updated widths in {REPORT}")


if __name__ == "__main__":
    main()
