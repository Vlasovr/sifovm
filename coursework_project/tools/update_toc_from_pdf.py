from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "documents" / "Пояснительная_записка.docx"
PDF = ROOT / "documents" / "Пояснительная_записка.pdf"
PDFINFO = Path(
    r"C:\Users\user\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
    r"\poppler-25.07.0\Library\bin\pdfinfo.exe"
)
PDFTOTEXT = Path(
    r"C:\Users\user\AppData\Local\Microsoft\WinGet\Packages\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
    r"\poppler-25.07.0\Library\bin\pdftotext.exe"
)


def normalize(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def page_count(pdf: Path) -> int:
    info = subprocess.check_output([str(PDFINFO), str(pdf)], text=True, encoding="utf-8", errors="ignore")
    match = re.search(r"Pages:\s+(\d+)", info)
    if not match:
        raise RuntimeError("Cannot determine PDF page count")
    return int(match.group(1))


def extract_pages(pdf: Path) -> dict[int, str]:
    pages = {}
    for page in range(1, page_count(pdf) + 1):
        text = subprocess.check_output(
            [str(PDFTOTEXT), "-layout", "-f", str(page), "-l", str(page), str(pdf), "-"],
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
        pages[page] = normalize(text)
    return pages


def paragraph_set_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def toc_entries(document: Document):
    entries = []
    for idx, paragraph in enumerate(document.paragraphs[:120]):
        text = paragraph.text.strip()
        if not re.search(r"\t\d+$", text):
            continue
        title, _old_page = text.rsplit("\t", 1)
        entries.append((idx, title))
    return entries


def appendix_key(title: str) -> str | None:
    match = re.match(r"ПРИЛОЖЕНИЕ\s+([А-Я])\b", title)
    if not match:
        return None
    return f"ПРИЛОЖЕНИЕ {match.group(1)}"


def numbered_key(title: str) -> str | None:
    match = re.match(r"(\d+(?:\.\d+)*)\s+", title)
    if not match:
        return None
    return match.group(1)


def find_page(title: str, pages: dict[int, str]) -> int:
    app_key = appendix_key(title)
    num_key = numbered_key(title)
    title_l = title.lower()

    for page, text in pages.items():
        if page < 4:
            continue
        text_l = text.lower()
        if app_key and re.search(rf"\b{re.escape(app_key)}\b", text, re.IGNORECASE):
            return page
        if title_l in text_l:
            return page

    # Fallbacks for long TOC labels whose body heading uses an expanded name.
    aliases = {
        "3.7 Функциональное моделирование АЛУ": r"3\.7\s+Функциональное моделирование арифметико-логического",
        "3.10 Функциональное моделирование КПДП": r"3\.10\s+Функциональное моделирование контроллера прямого",
    }
    pattern = aliases.get(title)
    if pattern:
        for page, text in pages.items():
            if re.search(pattern, text, re.IGNORECASE):
                return page

    if num_key:
        suffix = title[len(num_key) :].strip()
        words = re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", suffix)
        prefix = words[: min(5, len(words))]
        if prefix:
            prefix_pattern = r"\s+".join(re.escape(word) for word in prefix)
            pattern = rf"(?<![\d.]){re.escape(num_key)}(?![\d.])\s+{prefix_pattern}"
            for page, text in pages.items():
                if page < 4:
                    continue
                if re.search(pattern, text, re.IGNORECASE):
                    return page

    raise RuntimeError(f"Heading not found in PDF: {title}")


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not PDF.exists():
        raise FileNotFoundError(PDF)

    document = Document(DOCX)
    pages = extract_pages(PDF)
    updates = []

    for idx, title in toc_entries(document):
        page = find_page(title, pages)
        paragraph_set_text(document.paragraphs[idx], f"{title}\t{page}")
        updates.append((idx, title, page))

    document.save(DOCX)
    for idx, title, page in updates:
        print(f"{idx}: {title} -> {page}")


if __name__ == "__main__":
    main()
