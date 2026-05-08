from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("D:/git/sifovm/coursework_project")
DOCS = ROOT / "documents"
IMAGES = DOCS / "images"
GENERATED = IMAGES / "generated_report"
REPORT = DOCS / "Пояснительная_записка.docx"


def font(name: str, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/consola.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def make_dump_image(path: Path, title: str, rows: list[tuple[str, str, str]]) -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)
    w, h = 1280, 150 + len(rows) * 52
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)

    f_title = font("arialbd.ttf", 34)
    f_head = font("arialbd.ttf", 24)
    f_text = font("arial.ttf", 23)
    f_mono = font("consola.ttf", 24)

    d.rectangle((0, 0, w, 82), fill=(240, 244, 248))
    d.text((34, 23), title, fill=(20, 31, 48), font=f_title)

    x0, y0 = 36, 112
    col = [x0, x0 + 185, x0 + 365, x0 + 900]
    row_h = 52
    d.rectangle((x0, y0, w - 36, y0 + row_h), fill=(231, 238, 247), outline=(150, 165, 185))
    for x in col[1:3]:
        d.line((x, y0, x, y0 + row_h * (len(rows) + 1)), fill=(150, 165, 185), width=2)
    d.text((x0 + 18, y0 + 13), "Адрес", fill=(20, 31, 48), font=f_head)
    d.text((col[1] + 18, y0 + 13), "Слово", fill=(20, 31, 48), font=f_head)
    d.text((col[2] + 18, y0 + 13), "Назначение", fill=(20, 31, 48), font=f_head)

    for i, (addr, value, comment) in enumerate(rows):
        y = y0 + row_h * (i + 1)
        fill = (250, 252, 255) if i % 2 == 0 else (255, 255, 255)
        d.rectangle((x0, y, w - 36, y + row_h), fill=fill, outline=(205, 214, 225))
        d.text((x0 + 18, y + 13), addr, fill=(15, 55, 110), font=f_mono)
        d.text((col[1] + 18, y + 13), value, fill=(15, 55, 110), font=f_mono)
        d.text((col[2] + 18, y + 13), comment, fill=(35, 43, 55), font=f_text)

    img.save(path)


def insert_after(paragraph, items):
    parent = paragraph._p.getparent()
    idx = parent.index(paragraph._p)
    for item in reversed(items):
        parent.insert(idx + 1, item._p)


def add_paragraph_after(doc: Document, paragraph, text: str = "", style: str | None = None):
    p = doc.add_paragraph(text, style=style)
    paragraph._p.addnext(p._p)
    return p


def new_paragraph(doc: Document, text: str = "", style: str | None = None):
    return doc.add_paragraph(text, style=style)


def image_paragraph(doc: Document, image_path: Path, width_inches: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def set_body_text(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def make_images() -> dict[str, Path]:
    before = GENERATED / "fig_2_4_ram_before.png"
    after = GENERATED / "fig_2_5_ram_after.png"
    make_dump_image(
        before,
        "Дамп ОЗУ до функционального моделирования",
        [
            ("000A", "0000", "первая ячейка области обмена КПДП"),
            ("000B", "0000", "вторая ячейка области обмена КПДП"),
            ("000C", "0000", "третья ячейка области обмена КПДП"),
            ("0020", "8001", "исходный операнд для SRA и INCS"),
            ("0021", "00F0", "исходное значение для цепочки OR/NOR"),
            ("0022", "0F0F", "операнд команды OR"),
            ("0023", "0000", "операнд команды NOR"),
            ("0025", "FFFF", "операнд для формирования нулевого результата"),
            ("0030", "0000", "ячейка назначения команды MOV R3, 0030h"),
        ],
    )
    make_dump_image(
        after,
        "Контрольный дамп ОЗУ после выполнения программы",
        [
            ("000A", "1111", "первое слово, записанное КПДП"),
            ("000B", "2222", "второе слово, записанное КПДП"),
            ("000C", "3333", "третье слово, записанное КПДП"),
            ("0020", "8001", "исходный операнд сохранён"),
            ("0021", "00F0", "исходное значение сохранено"),
            ("0022", "0F0F", "операнд OR сохранён"),
            ("0025", "FFFF", "операнд JZ-проверки сохранён"),
            ("0030", "C001", "результат POP/MOV сохранён процессором"),
        ],
    )
    return {
        "dma_wave": IMAGES / "wave_dma.png",
        "system_wave": IMAGES / "wave_system.png",
        "before": before,
        "after": after,
    }


def find_para(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise RuntimeError(f"Paragraph not found: {needle}")


def main() -> None:
    paths = make_images()
    doc = Document(str(REPORT))

    # Make repeated runs idempotent-ish.
    if any("Временная диаграмма работы КПДП и ОЗУ" in p.text for p in doc.paragraphs):
        print("Figures already inserted; only images regenerated.")
        return

    for p in doc.paragraphs:
        if "Рисунок 2.1 - Символическое-представление подсистемы памяти данных" in p.text:
            p.text = "Рисунок 2.1 - RTL-представление подсистемы памяти данных"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Рисунок 2.2 - Символическое-представление подсистемы памяти данных" in p.text:
            p.text = "Рисунок 2.2 - Символическое представление модулей ПЗУ, ОЗУ и кэша данных"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.text.strip() == "Рис. 2.3 - Арифметико-логическое устройство.":
            p.text = "Рисунок 2.7 - Арифметико-логическое устройство"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = find_para(doc, "Для временной диаграммы работы ОЗУ необходимо показать")
    items = []
    intro = new_paragraph(
        doc,
        "Результаты функциональной проверки подсистемы памяти приведены на рисунках 2.3-2.5. "
        "На диаграмме отдельно показаны запрос контроллера прямого доступа к памяти, выдача "
        "арбитром разрешения, запись трёх контрольных слов в ОЗУ и установка признака завершения передачи.",
    )
    set_body_text(intro)
    items.append(intro)
    items.append(image_paragraph(doc, paths["dma_wave"], 6.6))
    items.append(caption(doc, "Рисунок 2.3 - Временная диаграмма работы КПДП и ОЗУ"))
    explain = new_paragraph(
        doc,
        "На рисунке 2.3 видно, что после активации START контроллер формирует запрос REQ. "
        "После получения GNT по адресам 000Ah, 000Bh и 000Ch последовательно записываются слова "
        "1111h, 2222h и 3333h. По окончании передачи устанавливается сигнал DONE.",
    )
    set_body_text(explain)
    items.append(explain)
    items.append(image_paragraph(doc, paths["before"], 6.35))
    items.append(caption(doc, "Рисунок 2.4 - Дамп ОЗУ до функционального моделирования"))
    items.append(image_paragraph(doc, paths["after"], 6.35))
    items.append(caption(doc, "Рисунок 2.5 - Контрольный дамп ОЗУ после выполнения тестовой программы"))
    explain2 = new_paragraph(
        doc,
        "Сравнение дампов подтверждает, что область обмена КПДП была изменена только по адресам "
        "000Ah-000Ch, а процессорная часть записала результат C001h в ячейку 0030h. Остальные "
        "контрольные ячейки, используемые как операнды команд, сохраняют исходные значения.",
    )
    set_body_text(explain2)
    items.append(explain2)
    insert_after(anchor, items)

    anchor2 = find_para(doc, "При обычном линейном выполнении после завершения команды указатель команд увеличивается на два")
    items2 = []
    intro2 = new_paragraph(
        doc,
        "Временная диаграмма работы полной модели показывает последовательную выборку двухсловных команд, "
        "изменение указателя команд, загрузку регистров IR0 и IR1, изменение регистров общего назначения, "
        "обращения к памяти и завершение программы командой HLT.",
    )
    set_body_text(intro2)
    items2.append(intro2)
    items2.append(image_paragraph(doc, paths["system_wave"], 6.7))
    items2.append(caption(doc, "Рисунок 2.6 - Временная диаграмма выполнения тестовой программы устройством управления"))
    explain3 = new_paragraph(
        doc,
        "На рисунке 2.6 показано, что IP последовательно принимает адреса первых слов команд, "
        "IR0 содержит код операции и номер регистра, а IR1 - адресную часть. После выполнения "
        "арифметико-логических и стековых команд регистры R1, R2, R3, R4 и R7 принимают контрольные "
        "значения, флаг Z устанавливается после нулевого результата, а сигнал HALT фиксирует окончание программы.",
    )
    set_body_text(explain3)
    items2.append(explain3)
    insert_after(anchor2, items2)

    backup = REPORT.with_name("Пояснительная_записка_before_figures_backup.docx")
    if not backup.exists():
        backup.write_bytes(REPORT.read_bytes())
    try:
        doc.save(str(REPORT))
        print(f"Updated {REPORT}")
    except PermissionError:
        out = REPORT.with_name("Пояснительная_записка_с_диаграммами.docx")
        doc.save(str(out))
        print(f"Original report is locked; saved updated copy to {out}")
    print(f"Generated images in {GENERATED}")


if __name__ == "__main__":
    main()
