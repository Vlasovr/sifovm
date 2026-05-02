from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import subprocess
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT = Path(__file__).resolve().parents[1]
DOCS = PROJECT / "documents"
IMAGES = DOCS / "images"
GRAPHICS = DOCS / "graphics"
SIM = PROJECT / "sim" / "ghdl"
PDFTOPPM = Path(
    r"C:\Users\user\AppData\Local\Microsoft\WinGet\Packages"
    r"\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
    r"\poppler-25.07.0\Library\bin\pdftoppm.exe"
)


PROGRAM_ROWS = [
    ("0000h", "0110h", "0020h", "MOV 0020h, R1", "загрузка отрицательного числа для проверки SRA"),
    ("0002h", "0510h", "0000h", "SRA R1", "арифметический сдвиг вправо, ожидается C000h и C=1"),
    ("0004h", "0610h", "0000h", "INCS R1", "инкремент на значение флага S, ожидается C001h"),
    ("0006h", "0710h", "0000h", "PUSH R1", "помещение C001h в стек"),
    ("0008h", "0120h", "0021h", "MOV 0021h, R2", "загрузка первого логического операнда"),
    ("000Ah", "0320h", "0022h", "OR R2, 0022h", "формирование 0FFFh"),
    ("000Ch", "0420h", "0023h", "NOR R2, 0023h", "формирование F000h"),
    ("000Eh", "0830h", "0000h", "POP R3", "извлечение C001h из стека"),
    ("0010h", "0230h", "0030h", "MOV R3, 0030h", "запись результата в ОЗУ"),
    ("0012h", "0140h", "0025h", "MOV 0025h, R4", "загрузка FFFFh для получения нуля"),
    ("0014h", "0440h", "0025h", "NOR R4, 0025h", "получение 0000h и установка Z=1"),
    ("0016h", "0A00h", "001Ah", "JZ 001Ah", "условный переход должен выполниться"),
    ("0018h", "0150h", "0024h", "MOV 0024h, R5", "команда должна быть пропущена"),
    ("001Ah", "0900h", "001Eh", "JMP 001Eh", "безусловный переход к чтению результата DMA"),
    ("001Ch", "0160h", "0024h", "MOV 0024h, R6", "команда должна быть пропущена"),
    ("001Eh", "0170h", "000Ah", "MOV 000Ah, R7", "чтение первого слова, записанного КПДП"),
    ("0020h", "0000h", "0000h", "HLT", "останов полной модели"),
]

RAM_ROWS = [
    ("000Ah", "0000h", "1111h", "первое слово, переданное КПДП"),
    ("000Bh", "0000h", "2222h", "второе слово, переданное КПДП"),
    ("000Ch", "0000h", "3333h", "третье слово, переданное КПДП"),
    ("0020h", "8001h", "8001h", "исходное отрицательное число для SRA"),
    ("0021h", "00F0h", "00F0h", "первый логический операнд"),
    ("0022h", "0F0Fh", "0F0Fh", "второй операнд для OR"),
    ("0023h", "0000h", "0000h", "операнд для NOR"),
    ("0024h", "1234h", "1234h", "контрольное значение пропущенных команд"),
    ("0025h", "FFFFh", "FFFFh", "операнд для получения нулевого результата"),
    ("0030h", "0000h", "C001h", "результат, записанный командой MOV R3, 0030h"),
]

TEST_ROWS = [
    ("АЛУ", "tb_alu_core", "SRA, INCS, OR, NOR; флаги Z/S/C/O", "tb_alu_core: TEST PASSED"),
    ("Стек", "tb_stack7x16", "PUSH, POP, SP, empty/full", "tb_stack7x16: TEST PASSED"),
    ("Память и КПДП", "tb_mem_dma", "REQ/GNT, запись 1111h/2222h/3333h по адресам 000Ah..000Ch", "tb_mem_dma: TEST PASSED"),
    ("Полная система", "tb_system", "программа, переходы, стек, кэш, КПДП, предсказатель", "tb_system: TEST PASSED"),
]

GRAPHIC_SHEETS = [
    ("Лист А1", "Лист_А1_Общая_структурная_схема.pdf", "Общая структурная схема микро-ЭВМ"),
    ("Лист А2", "Лист_А2_Устройство_управления.pdf", "Устройство управления"),
    ("Лист А3", "Лист_А3_Исполнительный_тракт.pdf", "Исполнительный тракт, АЛУ и стек"),
    ("Лист А4", "Лист_А4_Память_DMA_КПДП_арбитр.pdf", "Память, кэш, КПДП и арбитр"),
    ("Доп. лист 1", "Доп_лист_1_Процессорное_ядро.pdf", "Процессорное ядро"),
    ("Доп. лист 2", "Доп_лист_2_Сигналы_управления.pdf", "Сигналы управления"),
    ("Доп. лист 3", "Доп_лист_3_Тракт_PC_IR.pdf", "Тракт PC/IR"),
    ("Доп. лист 4", "Доп_лист_4_Регистровый_файл.pdf", "Регистровый файл"),
    ("Доп. лист 5", "Доп_лист_5_АЛУ_и_флаги.pdf", "АЛУ и регистр флагов"),
    ("Доп. лист 6", "Доп_лист_6_Стек.pdf", "Стек"),
    ("Доп. лист 7", "Доп_лист_7_ПЗУ_команд.pdf", "ПЗУ команд"),
    ("Доп. лист 8", "Доп_лист_8_ОЗУ_данных.pdf", "ОЗУ данных"),
    ("Доп. лист 9", "Доп_лист_9_Кэш_данных.pdf", "Кэш данных"),
    ("Доп. лист 10", "Доп_лист_10_КПДП.pdf", "Контроллер прямого доступа к памяти"),
    ("Доп. лист 11", "Доп_лист_11_Арбитр_шины.pdf", "Арбитр шины"),
    ("Доп. лист 12", "Доп_лист_12_Предсказатель_A4.pdf", "Предсказатель переходов A4"),
]


@dataclass
class VcdData:
    vars: dict[str, str]
    changes: dict[str, list[tuple[float, str]]]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_doc(doc: Document, landscape: bool = False) -> None:
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(14)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[name].font.color.rgb = RGBColor(0, 0, 0)
        styles[name].font.bold = True


def add_title(doc: Document, title: str, subtitle: str = "") -> None:
    for text, size, bold in [
        ("Министерство образования Республики Беларусь", 12, False),
        ("Учреждение образования", 12, False),
        ("Белорусский государственный университет информатики и радиоэлектроники", 12, False),
        ("Кафедра электронных вычислительных машин", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(14)
    for _ in range(8):
        doc.add_paragraph()
    for line in [
        "Студент: Власов",
        "Руководитель: ____________________",
        "Минск 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if "Студент" in line or "Руководитель" in line else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(14)
    doc.add_page_break()


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = widths is None
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(12)


def add_picture(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def parse_vcd(path: Path) -> VcdData:
    scope: list[str] = []
    vars_by_id: dict[str, str] = {}
    changes: dict[str, list[tuple[float, str]]] = {}
    current_time = 0.0
    in_defs = True
    for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("$scope"):
            parts = line.split()
            scope.append(parts[2])
            continue
        if line.startswith("$upscope"):
            if scope:
                scope.pop()
            continue
        if line.startswith("$var"):
            parts = line.split()
            ident = parts[3]
            name = parts[4]
            full = ".".join(scope + [name])
            vars_by_id[ident] = full
            changes.setdefault(full, [])
            continue
        if line.startswith("$enddefinitions"):
            in_defs = False
            continue
        if in_defs:
            continue
        if line.startswith("#"):
            current_time = int(line[1:]) / 1_000_000.0
            continue
        if line[0] in "01xXzZuU" and len(line) >= 2:
            ident = line[1:]
            value = line[0].lower()
            full = vars_by_id.get(ident)
            if full:
                changes[full].append((current_time, value))
            continue
        if line.startswith("b"):
            value, ident = line.split(maxsplit=1)
            full = vars_by_id.get(ident)
            if full:
                changes[full].append((current_time, value[1:].lower()))
    return VcdData(vars_by_id, changes)


def select_changes(data: VcdData, top: str, name: str) -> list[tuple[float, str]]:
    candidates = [f"{top}.{name}", name]
    for candidate in candidates:
        if candidate in data.changes:
            return data.changes[candidate]
    for full, changes in data.changes.items():
        if full.endswith("." + name):
            return changes
    return []


def format_value(value: str) -> str:
    if not value:
        return ""
    if len(value) == 1 and value in "01xzui":
        return value.upper()
    if set(value) <= {"0", "1"}:
        return f"{int(value, 2):0{max(1, (len(value)+3)//4)}X}h"
    return value.upper()


def draw_waveform(
    vcd: Path,
    top: str,
    rows: list[tuple[str, str]],
    start_ns: float,
    end_ns: float,
    out: Path,
    title: str,
) -> None:
    data = parse_vcd(vcd)
    row_h = 42
    left = 210
    right = 40
    top_margin = 70
    width = 1800
    height = top_margin + row_h * len(rows) + 60
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = load_font(18, bold=False)
    small = load_font(14, bold=False)
    bold = load_font(20, bold=True)
    draw.text((20, 18), title, fill=(20, 35, 55), font=bold)
    plot_w = width - left - right

    def x_of(t: float) -> int:
        return int(left + (t - start_ns) / (end_ns - start_ns) * plot_w)

    for i in range(11):
        t = start_ns + (end_ns - start_ns) * i / 10
        x = x_of(t)
        draw.line((x, top_margin - 10, x, height - 45), fill=(226, 232, 240))
        draw.text((x - 22, height - 38), f"{t:.0f} ns", fill=(75, 85, 99), font=small)

    for idx, (label, name) in enumerate(rows):
        y = top_margin + idx * row_h
        draw.line((15, y + row_h - 4, width - 20, y + row_h - 4), fill=(235, 238, 242))
        draw.text((18, y + 11), label, fill=(15, 23, 42), font=font)
        changes = select_changes(data, top, name)
        draw_signal(draw, changes, start_ns, end_ns, x_of, y, row_h, left, width - right, font, small)
    image.save(out)


def draw_signal(draw, changes, start_ns, end_ns, x_of, y, row_h, x0, x1, font, small) -> None:
    if not changes:
        draw.line((x0, y + row_h // 2, x1, y + row_h // 2), fill=(180, 180, 180), width=2)
        return
    points = sorted(changes)
    prev = points[0][1]
    for t, value in points:
        if t <= start_ns:
            prev = value
        else:
            break
    visible = [(start_ns, prev)] + [(t, v) for t, v in points if start_ns < t <= end_ns] + [(end_ns, None)]
    is_scalar = all(len(v) == 1 for _, v in visible if v is not None)
    if is_scalar:
        hi = y + 10
        lo = y + row_h - 12
        last_x = x_of(start_ns)
        last_y = hi if prev == "1" else lo
        for t, value in visible[1:]:
            x = x_of(t)
            draw.line((last_x, last_y, x, last_y), fill=(37, 99, 235), width=2)
            if value is not None:
                new_y = hi if value == "1" else lo
                draw.line((x, last_y, x, new_y), fill=(37, 99, 235), width=2)
                last_y = new_y
                last_x = x
        return

    mid = y + row_h // 2
    for (t0, value), (t1, _) in zip(visible[:-1], visible[1:]):
        x_start = x_of(t0)
        x_end = x_of(t1)
        draw.rounded_rectangle(
            (x_start + 2, y + 8, max(x_start + 8, x_end - 2), y + row_h - 10),
            radius=4,
            outline=(59, 130, 246),
            fill=(239, 246, 255),
            width=1,
        )
        text = format_value(value)
        if x_end - x_start > 42:
            draw.text((x_start + 6, mid - 8), text, fill=(30, 64, 175), font=small)


def load_font(size: int, bold: bool = False):
    names = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            pass
    return ImageFont.load_default()


def render_pdf_sheets() -> list[Path]:
    IMAGES.mkdir(parents=True, exist_ok=True)
    rendered: list[Path] = []
    for index, (_, pdf_name, _) in enumerate(GRAPHIC_SHEETS, 1):
        pdf = GRAPHICS / pdf_name
        prefix = IMAGES / f"scheme_{index:02d}"
        subprocess.run(
            [str(PDFTOPPM), "-png", "-r", "120", "-singlefile", str(pdf), str(prefix)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        rendered.append(prefix.with_suffix(".png"))
    return rendered


def generate_waveforms() -> dict[str, Path]:
    IMAGES.mkdir(parents=True, exist_ok=True)
    outputs = {
        "alu": IMAGES / "wave_alu.png",
        "stack": IMAGES / "wave_stack.png",
        "dma": IMAGES / "wave_dma.png",
        "system": IMAGES / "wave_system.png",
    }
    draw_waveform(
        SIM / "tb_alu_core.vcd",
        "tb_alu_core",
        [("A", "a[15:0]"), ("B", "b[15:0]"), ("S_FLAG", "sflag"), ("OP", "op[2:0]"), ("Y", "y[15:0]"), ("Z", "z"), ("S", "s"), ("C", "c"), ("O", "o")],
        0,
        60,
        outputs["alu"],
        "Временная диаграмма АЛУ",
    )
    draw_waveform(
        SIM / "tb_stack7x16.vcd",
        "tb_stack7x16",
        [("CLK", "clk"), ("RESET", "rst"), ("PUSH", "push"), ("POP", "pop"), ("DIN", "din[15:0]"), ("DOUT", "dout[15:0]"), ("SP", "sp[2:0]"), ("EMPTY", "empty"), ("FULL", "full")],
        0,
        130,
        outputs["stack"],
        "Временная диаграмма стека",
    )
    draw_waveform(
        SIM / "tb_mem_dma.vcd",
        "tb_mem_dma",
        [("CLK", "clk"), ("START", "start"), ("DEV_VALID", "dev_valid"), ("DEV_DATA", "dev_data[15:0]"), ("REQ", "dma_req"), ("GNT", "dma_gnt"), ("RAM_WE", "ram_we"), ("RAM_ADDR", "ram_addr[15:0]"), ("RAM_DIN", "ram_din[15:0]"), ("DONE", "dma_done")],
        0,
        120,
        outputs["dma"],
        "Временная диаграмма КПДП и ОЗУ",
    )
    draw_waveform(
        SIM / "tb_system.vcd",
        "tb_system",
        [("CLK", "clk"), ("RESET", "rst"), ("DMA_START", "dma_start"), ("DMA_VALID", "dma_valid"), ("DMA_DATA", "dma_data[15:0]"), ("STATE", "dbg_state[4:0]"), ("PC", "dbg_pc[15:0]"), ("IR0", "dbg_ir0[15:0]"), ("IR1", "dbg_ir1[15:0]"), ("R1", "dbg_r1[15:0]"), ("R2", "dbg_r2[15:0]"), ("R3", "dbg_r3[15:0]"), ("R4", "dbg_r4[15:0]"), ("R7", "dbg_r7[15:0]"), ("FLAGS", "dbg_flags[3:0]"), ("SP", "dbg_sp[2:0]"), ("REQ_CPU", "req_cpu"), ("REQ_DMA", "req_dma"), ("GNT_CPU", "gnt_cpu"), ("GNT_DMA", "gnt_dma"), ("RAM_WE", "ram_we"), ("RAM_ADDR", "ram_addr[15:0]"), ("BP_HIST", "bp_hist[1:0]"), ("BP_PRED", "bp_pred"), ("HALT", "halt"), ("DMA_DONE", "dma_done")],
        0,
        1800,
        outputs["system"],
        "Временная диаграмма полной системы",
    )
    return outputs


def build_report(scheme_images: list[Path], waves: dict[str, Path]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", "к курсовому проекту на тему «Разработка микро-ЭВМ на ПЛИС»")
    doc.add_heading("СОДЕРЖАНИЕ", 1)
    for line in [
        "Введение",
        "1 Разработка общей структуры микро-ЭВМ",
        "2 Разработка основных устройств микро-ЭВМ",
        "3 Функциональное моделирование",
        "4 Анализ и оптимизация",
        "Заключение",
        "Список использованных источников",
        "Приложения",
    ]:
        add_body(doc, line)
    doc.add_page_break()

    doc.add_heading("ВВЕДЕНИЕ", 1)
    add_body(doc, "Целью курсового проекта является разработка учебной микро-ЭВМ на ПЛИС в соответствии с индивидуальным заданием по дисциплине «Структурная и функциональная организация ЭВМ». Система построена по Гарвардской архитектуре: память команд и память данных разделены, что упрощает тракт выборки команд и обмен с ОЗУ.")
    add_body(doc, "В проект включены 12 регистров общего назначения, 16-разрядное АЛУ, регистр флагов, стек глубиной 7 слов с ростом вниз, синхронные ПЗУ и ОЗУ, кэш данных, контроллер прямого доступа к памяти, арбитр шины и предсказатель переходов A4.")
    add_body(doc, "Практическая проверка выполнена в Quartus II 9.1 и GHDL. Полная компиляция проекта завершается без ошибок, а тестовые стенды подтверждают корректность работы основных узлов.")

    doc.add_heading("1 РАЗРАБОТКА ОБЩЕЙ СТРУКТУРЫ МИКРО-ЭВМ", 1)
    doc.add_heading("1.1 Функциональный состав", 2)
    add_body(doc, "Микро-ЭВМ состоит из процессорного ядра, подсистемы памяти и внешнего канала обмена через КПДП. Верхний модуль проекта называется microcomputer_top, а внутренняя структурная схема вынесена в system_core.")
    add_table(doc, ["Блок", "Назначение"], [
        ("ПЗУ команд", "Хранение двухсловных команд программы и выдача слова по адресу PC."),
        ("ОЗУ данных", "Хранение операндов, результатов и области обмена с КПДП."),
        ("Регистровый файл 12 x 16", "Быстрое хранение промежуточных данных процессора."),
        ("АЛУ", "Выполнение операций INCS, OR, NOR и SRA с формированием флагов Z/S/C/O."),
        ("Стек 7 x 16", "Временное хранение данных при PUSH и POP, рост вниз."),
        ("Кэш данных", "4-way кэш со сквозной записью и замещением по наибольшей давности хранения."),
        ("КПДП", "Запись трёх 16-разрядных слов в ОЗУ начиная с адреса 000Ah."),
        ("Арбитр", "Централизованное параллельное распределение доступа к ОЗУ между CPU и КПДП."),
        ("Предсказатель A4", "Прогноз условных переходов по индексу PC(2)||GHR(2)."),
    ], [4.5, 11.5])
    add_picture(doc, scheme_images[0], 15.5, "Рисунок 1.1 - Общая структурная схема микро-ЭВМ")

    doc.add_heading("1.2 Система команд", 2)
    add_body(doc, "Принят фиксированный двухсловный формат команды. Первое слово содержит код операции и номер регистра, второе слово используется как прямой адрес операнда, адрес перехода или резервное поле.")
    add_table(doc, ["Слово", "Разряды", "Содержимое"], [
        ("Word0", "15..8", "код операции OPCODE"),
        ("Word0", "7..4", "номер регистра Rn"),
        ("Word0", "3..0", "резерв"),
        ("Word1", "15..0", "адрес операнда или адрес перехода"),
    ], [3, 4, 9])
    add_table(doc, ["Код", "Мнемоника", "Действие"], [
        ("00h", "HLT", "останов выполнения"),
        ("01h", "MOV adr, Rn", "чтение слова из ОЗУ в регистр"),
        ("02h", "MOV Rn, adr", "запись регистра в ОЗУ"),
        ("03h", "OR Rn, adr", "логическое ИЛИ регистра и слова памяти"),
        ("04h", "NOR Rn, adr", "инверсия OR"),
        ("05h", "SRA Rn", "арифметический сдвиг вправо"),
        ("06h", "INCS Rn", "инкремент на значение флага S"),
        ("07h", "PUSH Rn", "запись регистра в стек"),
        ("08h", "POP Rn", "чтение слова из стека в регистр"),
        ("09h", "JMP adr", "безусловный переход"),
        ("0Ah", "JZ adr", "условный переход при Z=1"),
    ], [2.2, 4, 9.8])

    doc.add_heading("2 РАЗРАБОТКА ОСНОВНЫХ УСТРОЙСТВ МИКРО-ЭВМ", 1)
    for title, text in [
        ("2.1 Запоминающие устройства", "ПЗУ команд и ОЗУ данных реализованы как синхронные словные устройства. Команда занимает два 16-разрядных слова, поэтому блок выборки последовательно получает IR0 и IR1."),
        ("2.2 Устройство управления", "Устройство управления организует фазы T0...T5 и состояние ожидания Tw при промахе кэша или занятой шине. Управляющие сигналы формируют чтение ПЗУ, обращение к кэшу, запись РОН, изменение SP и обновление предсказателя."),
        ("2.3 Арифметико-логическое устройство", "АЛУ поддерживает четыре обязательные операции: OR, NOR, SRA и INCS. После выполнения обновляются флаги Z, S, C и O."),
        ("2.4 Кэш, КПДП и арбитр", "Кэш данных имеет четыре направления и шестнадцать множеств. КПДП получает приоритет в арбитре и записывает входной поток в ОЗУ по адресам 000Ah, 000Bh и 000Ch."),
        ("2.5 Предсказатель переходов", "Предсказатель A4 использует двухбитовую глобальную историю GHR и младшие биты PC для обращения к таблицам PHT и BTB."),
    ]:
        doc.add_heading(title, 2)
        add_body(doc, text)
    add_picture(doc, scheme_images[1], 15.5, "Рисунок 2.1 - Устройство управления")
    add_picture(doc, scheme_images[2], 15.5, "Рисунок 2.2 - Исполнительный тракт")
    add_picture(doc, scheme_images[3], 15.5, "Рисунок 2.3 - Память, КПДП и арбитр")

    doc.add_heading("3 ФУНКЦИОНАЛЬНОЕ МОДЕЛИРОВАНИЕ", 1)
    add_body(doc, "Для проверки разработаны четыре тестовых стенда. Они выполняются автоматически и завершаются диагностическим сообщением TEST PASSED.")
    add_table(doc, ["Проверка", "Testbench", "Контролируемые сигналы", "Итог"], TEST_ROWS, [3, 3.5, 6.5, 3.5])
    add_picture(doc, waves["alu"], 16, "Рисунок 3.1 - Проверка операций АЛУ")
    add_picture(doc, waves["stack"], 16, "Рисунок 3.2 - Проверка стека")
    add_picture(doc, waves["dma"], 16, "Рисунок 3.3 - Проверка КПДП и записи в ОЗУ")
    add_picture(doc, waves["system"], 16, "Рисунок 3.4 - Проверка полной системы")
    add_table(doc, ["Адрес", "До запуска", "После HLT", "Пояснение"], RAM_ROWS, [2.5, 3, 3, 7])

    doc.add_heading("4 АНАЛИЗ И ОПТИМИЗАЦИЯ", 1)
    add_body(doc, "По результатам компиляции в Quartus II 9.1 проект собирается успешно. Верхний модуль содержит только контекстно значимые внешние порты: тактовый вход, сброс, интерфейс входного потока КПДП и два выходных признака завершения.")
    add_table(doc, ["Показатель", "Значение"], [
        ("Top-level entity", "microcomputer_top"),
        ("Семейство ПЛИС", "Cyclone II"),
        ("Результат Quartus", "Full Compilation was successful, 0 errors"),
        ("Результат полной симуляции", "tb_system: TEST PASSED на 1675 ns"),
    ], [5, 10])
    add_body(doc, "Физические номера выводов не фиксировались, так как в задании не указана конкретная учебная плата. При переносе на плату необходимо назначить CLK, RESET, DMA_START, DMA_VALID, DMA_DATA[15:0], HALT и DMA_DONE согласно её pinout.")

    doc.add_heading("ЗАКЛЮЧЕНИЕ", 1)
    add_body(doc, "В ходе курсового проекта разработана микро-ЭВМ на ПЛИС с 16-разрядными шинами адреса и данных, раздельной памятью команд и данных, регистровым файлом 12 x 16, АЛУ, стеком, кэшем, КПДП, арбитром и предсказателем переходов A4.")
    add_body(doc, "Подготовлены VHDL-описания, проект Quartus II, тестовые стенды, графические листы и приложения. Компиляция и моделирование подтверждают корректность работы контрольной программы.")

    doc.add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1)
    for item in [
        "Методические материалы по дисциплине «Структурная и функциональная организация ЭВМ».",
        "Altera Quartus II 9.1 Web Edition. Руководство пользователя.",
        "IEEE Standard VHDL Language Reference Manual.",
        "Материалы 1.BO6, 2.ALU6, 3.CUU6, 4.MPUU6, 5.ADDRESS из учебного комплекта.",
    ]:
        add_body(doc, item)

    doc.add_heading("ПРИЛОЖЕНИЯ", 1)
    for name in [
        "Приложение А - графическая часть.",
        "Приложение Б - листинг программы и дампы памяти.",
        "Приложение В - VHDL-листинги.",
        "Приложение Г - материалы тестирования.",
        "Приложение Д - инструкция по снятию временных диаграмм.",
    ]:
        add_body(doc, name)

    out = DOCS / "Пояснительная_записка.docx"
    doc.save(out)
    return out


def build_appendix_a(scheme_images: list[Path]) -> Path:
    doc = Document()
    setup_doc(doc, landscape=True)
    for index, ((sheet, _, title), image) in enumerate(zip(GRAPHIC_SHEETS, scheme_images), 1):
        if index == 1:
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_after = Pt(8)
            title_run = title_p.add_run("ПРИЛОЖЕНИЕ А. Графическая часть")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.name = "Times New Roman"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{sheet}. {title}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_after = Pt(0)
        img_p.add_run().add_picture(str(image), width=Cm(21.8))
    out = DOCS / "Приложение_А_Графическая_часть.docx"
    doc.save(out)
    return out


def build_appendix_b() -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "ПРИЛОЖЕНИЕ Б", "Листинг программы и дампы памяти")
    doc.add_heading("Б.1 Листинг контрольной программы", 1)
    add_table(doc, ["Адрес", "Слово 0", "Слово 1", "Команда", "Назначение"], PROGRAM_ROWS, [2, 2.2, 2.2, 3.5, 7])
    doc.add_heading("Б.2 Контрольные дампы ОЗУ", 1)
    add_table(doc, ["Адрес", "До запуска", "После HLT", "Пояснение"], RAM_ROWS, [2.5, 3, 3, 7])
    doc.add_heading("Б.3 MIF-файлы", 1)
    for mif in ["rom_program.mif", "ram_initial.mif", "ram_expected_after_hlt.mif"]:
        doc.add_heading(mif, 2)
        add_code_block(doc, (PROJECT / "memory" / mif).read_text(encoding="utf-8"))
    out = DOCS / "Приложение_Б_Листинг_и_дампы.docx"
    doc.save(out)
    return out


def build_appendix_c() -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "ПРИЛОЖЕНИЕ В", "VHDL-листинги")
    for source in sorted((PROJECT / "src").glob("*.vhd")) + sorted((PROJECT / "tb").glob("*.vhd")):
        doc.add_heading(source.relative_to(PROJECT).as_posix(), 1)
        add_code_block(doc, source.read_text(encoding="utf-8"))
    out = DOCS / "Приложение_В_VHDL_листинги.docx"
    doc.save(out)
    return out


def build_appendix_d(waves: dict[str, Path]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "ПРИЛОЖЕНИЕ Г", "Материалы тестирования")
    add_table(doc, ["Проверка", "Testbench", "Контролируемые сигналы", "Итог"], TEST_ROWS, [3, 3.5, 6.5, 3.5])
    for key, caption in [
        ("alu", "Г.1 Временная диаграмма АЛУ"),
        ("stack", "Г.2 Временная диаграмма стека"),
        ("dma", "Г.3 Временная диаграмма КПДП"),
        ("system", "Г.4 Временная диаграмма полной системы"),
    ]:
        add_picture(doc, waves[key], 16, caption)
    out = DOCS / "Приложение_Г_Тестирование.docx"
    doc.save(out)
    return out


def build_appendix_e() -> Path:
    doc = Document()
    setup_doc(doc)
    add_title(doc, "ПРИЛОЖЕНИЕ Д", "Инструкция по снятию временных диаграмм")
    for title, body in testing_instruction_sections():
        doc.add_heading(title, 1)
        for paragraph in body:
            add_body(doc, paragraph)
    out = DOCS / "Приложение_Д_Инструкция_по_диаграммам.docx"
    doc.save(out)
    write_testing_markdown()
    return out


def testing_instruction_sections() -> list[tuple[str, list[str]]]:
    return [
        ("1 Подготовка", [
            "Открой проект quartus/coursework.qpf в Quartus II 9.1 и убедись, что Processing -> Start Compilation завершается без ошибок.",
            "Для ModelSim используй скрипт quartus/run_tb_system.do. В transcript должна появиться строка tb_system: TEST PASSED.",
        ]),
        ("2 Диаграмма АЛУ", [
            "Запусти tb_alu_core. Добавь сигналы a, b, sflag, op, y, z, s, c, o. На диаграмме должны быть четыре участка: SRA для 8001h, INCS для C000h при S=1, OR для 00F0h и 0F0Fh, NOR для 0FFFh.",
            "Проверь, что после SRA формируется C000h и C=1, после INCS получается C001h, после OR получается 0FFFh, после NOR получается F000h.",
        ]),
        ("3 Диаграмма стека", [
            "Запусти tb_stack7x16. Добавь clk, rst, push, pop, din, dout, sp, empty, full. Сними участок, где выполняются две записи PUSH и последующее чтение POP.",
            "Проверь рост вниз: после PUSH указатель SP уменьшается, после POP увеличивается. После возврата SP к 7 стек считается пустым.",
        ]),
        ("4 Диаграмма КПДП и арбитра", [
            "Запусти tb_mem_dma или полный tb_system. Добавь start, dev_valid, dev_data, dma_req, dma_gnt, ram_we, ram_addr, ram_din, dma_done.",
            "На участке записи должны быть адреса 000Ah, 000Bh, 000Ch и данные 1111h, 2222h, 3333h. При одновременном запросе приоритет получает КПДП.",
        ]),
        ("5 Диаграмма полной системы", [
            "Запусти tb_system на 5 us. Добавь clk, rst, dma_start, dma_valid, dma_data, dbg_state, dbg_pc, dbg_ir0, dbg_ir1, dbg_r1, dbg_r2, dbg_r3, dbg_r4, dbg_r7, dbg_flags, dbg_sp, req_cpu, req_dma, gnt_cpu, gnt_dma, ram_we, ram_addr, bp_hist, bp_pred, halt, dma_done.",
            "Сними общий участок до HLT и отдельное приближение вокруг DMA-записи. В конце должны быть R1=C001h, R2=F000h, R3=C001h, R4=0000h, R7=1111h, FLAGS=1000, SP=7, DMA_DONE=1, BP_HIST=01.",
        ]),
    ]


def write_testing_markdown() -> None:
    lines = ["# Что снять на временных диаграммах", ""]
    for title, paragraphs in testing_instruction_sections():
        lines += [f"## {title}", ""]
        for paragraph in paragraphs:
            lines += [paragraph, ""]
    (DOCS / "Что_снимать_на_диаграммах.md").write_text("\n".join(lines), encoding="utf-8")


def add_code_block(doc: Document, code: str) -> None:
    for chunk in split_code(code, 2800):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(chunk)
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        r.font.size = Pt(8)


def split_code(code: str, limit: int) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    size = 0
    for line in code.splitlines():
        if size + len(line) + 1 > limit and current:
            chunks.append("\n".join(current))
            current = []
            size = 0
        current.append(line)
        size += len(line) + 1
    if current:
        chunks.append("\n".join(current))
    return chunks


def strip_variant_mentions_from_docx(path: Path) -> None:
    with ZipFile(path) as z:
        names = z.namelist()
        xml = z.read("word/document.xml").decode("utf-8")
    if re.search(r"variant|вариант", xml, re.IGNORECASE):
        raise RuntimeError(f"Unexpected variant mention in {path}")


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    IMAGES.mkdir(exist_ok=True)
    scheme_images = render_pdf_sheets()
    waves = generate_waveforms()
    outputs = [
        build_report(scheme_images, waves),
        build_appendix_a(scheme_images),
        build_appendix_b(),
        build_appendix_c(),
        build_appendix_d(waves),
        build_appendix_e(),
    ]
    for path in outputs:
        strip_variant_mentions_from_docx(path)
        print(path)


if __name__ == "__main__":
    main()
