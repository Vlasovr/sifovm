from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


REPORT = Path("D:/git/sifovm/coursework_project/documents/Пояснительная_записка.docx")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def has_picture(paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="pic"]'))


def style_cell(cell, width_cm: float | None = None) -> None:
    if width_cm is not None:
        cell.width = Cm(width_cm)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def polish_phase_table(doc: Document) -> None:
    rows = [
        ("Фаза", "Основные состояния RTL", "Назначение"),
        ("T0", "S_FETCH0_REQ", "выдача адреса IP в ПЗУ"),
        ("T1", "S_FETCH0_WAIT; S_FETCH0_LATCH", "ожидание и загрузка IR0"),
        ("T2", "S_FETCH1_REQ; S_FETCH1_WAIT; S_FETCH1_LATCH", "выборка и загрузка IR1"),
        ("T3", "S_DECODE", "декодирование КОП и номера регистра"),
        ("T4", "S_CACHE_READ; S_CACHE_WRITE; S_ALU_SETUP_REG; S_STACK_PUSH; S_STACK_POP; S_BRANCH_JMP; S_BRANCH_JZ", "выполнение операции"),
        ("Tw", "ожидание cache_ready_i или grant", "ожидание памяти, кэша или арбитра"),
        ("T5", "S_ALU_WRITE; S_RF_WRITE_MEM; S_RF_WRITE_POP; S_FINISH", "запись результата, обновление флагов и переход к следующей команде"),
    ]
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "S_FETCH0_REQ" not in text or "S_FETCH1_REQ" not in text:
            continue

        while len(table.rows) > len(rows):
            table._tbl.remove(table.rows[-1]._tr)
        while len(table.rows) < len(rows):
            table.add_row()

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [1.15, 9.1, 4.8]
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx]
            for col_idx, value in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = value
                style_cell(cell, widths[col_idx])
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        break


def main() -> None:
    doc = Document(str(REPORT))

    # Remove accidental uncaptioned top-level Quartus waveform screenshot left after Figure 2.2.
    to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if not has_picture(paragraph) or i == 0:
            continue
        prev = doc.paragraphs[i - 1].text.strip()
        nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
        if prev.startswith("Рисунок 2.2 -") and not nxt.startswith("Рисунок"):
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)

    polish_phase_table(doc)
    doc.save(str(REPORT))
    print(f"Polished {REPORT}")


if __name__ == "__main__":
    main()
