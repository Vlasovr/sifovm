from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


EMU_PER_IN = 914400
TWIPS_PER_IN = 1440


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "documents" / "Пояснительная_записка.docx"


FIGURES = [
    # (first image paragraph index, last caption paragraph index)
    (885, 886),
    (896, 897),
    (907, 908),
    (918, 919),
    (929, 932),
    (940, 943),
    (951, 954),
    (964, 966),
    (976, 978),
    (988, 990),
    (1000, 1002),
    (1012, 1014),
]


TEXT_REPLACEMENTS = {
    934: (
        "T0. После сброса указатель стека SP находится в состоянии 7h, "
        "соответствующем пустому стеку. На диаграмме контролируется именно "
        "изменение SP и разрешение записи результата в регистровый файл."
    ),
    935: (
        "T1. При выполнении команды PUSH R1 в стек заносится слово C001h, "
        "полученное после операций SRA и INCS. Указатель SP уменьшается с 7h "
        "до 6h, что соответствует занятию верхней ячейки стека."
    ),
    936: (
        "T2. При выполнении команды POP R3 сохранённое слово считывается из "
        "стека и записывается в регистр R3. Указатель SP увеличивается с 6h "
        "до 7h, поэтому стек снова переходит в пустое состояние."
    ),
    950: (
        "Результаты функционального моделирования арифметико-логического "
        "устройства показаны на рисунке 3.7. Проверяются операции, заданные "
        "индивидуальным вариантом: OR, NOR, SRA и INCS. На диаграмме "
        "контролируются код текущей команды, признаки выбранной операции, "
        "изменение регистров R1, R2, R4 и состояние регистра флагов."
    ),
    981: (
        "T1. При IR0 = 0A00h и IR1 = 001Ah устройство управления переходит "
        "к обработке условного перехода. Так как в регистре флагов установлен "
        "признак Z, значение PC заменяется адресом 001Ah."
    ),
    982: (
        "T2. После выполнения условного перехода обновляется история "
        "предсказателя BP_HIST. Изменение этого сигнала показывает, что "
        "фактический исход перехода был передан в блок предсказания."
    ),
    983: (
        "T3. Следующая команда JMP 001Eh является безусловной. Для неё "
        "направление не требуется предсказывать: адрес из IR1 напрямую "
        "загружается в PC, что видно по скачку указателя команд к 001Eh."
    ),
    992: (
        "T0. При активном dma_start контроллер переходит из состояния ожидания "
        "в состояние запроса шины. На вход dma_data подаётся первое слово "
        "1111h, а dma_valid подтверждает его действительность."
    ),
    993: (
        "T1. После получения разрешения gnt_dma контроллер активирует запись "
        "в ОЗУ. На адресную шину подаётся 000Ah, на шину данных - 1111h, "
        "сигнал ram_we устанавливается в единицу."
    ),
}


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def remove_section_properties(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for sect in list(p_pr.findall(qn("w:sectPr"))):
        p_pr.remove(sect)


def clone_section(base_sect, landscape: bool):
    sect = deepcopy(base_sect)

    # Section breaks for figure pages should start on a new page.
    for old_type in list(sect.findall(qn("w:type"))):
        sect.remove(old_type)
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect.insert(0, sect_type)

    pg_sz = sect.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect.append(pg_sz)

    if landscape:
        pg_sz.set(qn("w:w"), "16838")
        pg_sz.set(qn("w:h"), "11906")
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")
        if qn("w:orient") in pg_sz.attrib:
            del pg_sz.attrib[qn("w:orient")]

    pg_mar = sect.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect.append(pg_mar)

    if landscape:
        pg_mar.set(qn("w:top"), str(int(0.45 * TWIPS_PER_IN)))
        pg_mar.set(qn("w:bottom"), str(int(0.45 * TWIPS_PER_IN)))
        pg_mar.set(qn("w:left"), str(int(0.70 * TWIPS_PER_IN)))
        pg_mar.set(qn("w:right"), str(int(0.50 * TWIPS_PER_IN)))

    return sect


def set_section_break(paragraph, sect) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    remove_section_properties(paragraph)
    p_pr.append(deepcopy(sect))


def drawing_extents(drawing):
    extents = [e for e in drawing.xpath('.//*[local-name()="extent"]') if e.get("cx") and e.get("cy")]
    if not extents:
        return None
    cx = int(extents[0].get("cx"))
    cy = int(extents[0].get("cy"))
    return cx, cy, extents


def resize_drawing(drawing, width_in: float) -> None:
    data = drawing_extents(drawing)
    if not data:
        return
    cx, cy, extents = data
    ratio = cy / cx
    new_cx = int(width_in * EMU_PER_IN)
    new_cy = int(new_cx * ratio)
    for ext in extents:
        ext.set("cx", str(new_cx))
        ext.set("cy", str(new_cy))


def add_page_break_once(paragraph) -> None:
    existing = paragraph._p.xpath('.//w:br[@w:type="page"]')
    if existing:
        return
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def content_width_landscape_in() -> float:
    page_w = 16838 / TWIPS_PER_IN
    left = 0.70
    right = 0.50
    return page_w - left - right


def content_height_landscape_in() -> float:
    page_h = 11906 / TWIPS_PER_IN
    top = 0.45
    bottom = 0.45
    return page_h - top - bottom


def main() -> None:
    document = Document(DOCX)
    paragraphs = document.paragraphs
    base_sect = deepcopy(document.sections[0]._sectPr)
    portrait = clone_section(base_sect, landscape=False)
    landscape = clone_section(base_sect, landscape=True)

    for idx, text in TEXT_REPLACEMENTS.items():
        set_paragraph_text(paragraphs[idx], text)

    figure_page_width = content_width_landscape_in()
    figure_page_height = content_height_landscape_in()

    for first_img_idx, caption_idx in FIGURES:
        prev_idx = first_img_idx - 1
        while prev_idx > 0 and not paragraphs[prev_idx].text.strip() and not paragraphs[prev_idx]._p.xpath(".//w:drawing"):
            prev_idx -= 1

        set_section_break(paragraphs[prev_idx], portrait)
        set_section_break(paragraphs[caption_idx], landscape)

        image_paragraphs = []
        drawings = []
        for p in paragraphs[first_img_idx:caption_idx]:
            ds = p._p.xpath(".//w:drawing")
            if ds:
                image_paragraphs.append(p)
                drawings.extend(ds)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0

        caption = paragraphs[caption_idx]
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = 0
        caption.paragraph_format.space_after = 0

        if not drawings:
            continue

        ratios = []
        for dr in drawings:
            data = drawing_extents(dr)
            if data:
                cx, cy, _ = data
                ratios.append(cy / cx)

        if len(drawings) == 1:
            target_width = figure_page_width
        elif caption_idx in (932, 1014):
            # Dense two-part diagrams remain readable only when split across
            # two landscape pages. The caption stays under the second part.
            target_width = figure_page_width
            add_page_break_once(paragraphs[first_img_idx])
        else:
            # Leave enough vertical room for the caption on the same landscape page.
            available_height = figure_page_height - 0.90 - 0.10 * (len(drawings) - 1)
            target_width = min(figure_page_width, available_height / sum(ratios))

        for dr in drawings:
            resize_drawing(dr, target_width)

    # Keep the final trailing section in portrait orientation.
    body_sect = document._body._element.sectPr
    if body_sect is not None:
        new_body = clone_section(base_sect, landscape=False)
        body = document._body._element
        body.remove(body_sect)
        body.append(new_body)

    document.save(DOCX)
    print(f"Updated {DOCX}")


if __name__ == "__main__":
    main()
