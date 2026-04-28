from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "Курсовая" / "Записка Власов.docx"
OUT_DIR = ROOT / "variant4_coursework_project" / "report"
OUTPUT = OUT_DIR / "Записка_Власов_вариант4_исправленная.docx"


PROGRAM_ROWS = [
    ("0000h", "0110h", "0020h", "MOV 0020h, R1", "Загрузка отрицательного числа для проверки SRA."),
    ("0002h", "0510h", "0000h", "SRA R1", "Арифметический сдвиг вправо, ожидается C000h и C=1."),
    ("0004h", "0610h", "0000h", "INCS R1", "Инкремент на значение флага S, ожидается C001h."),
    ("0006h", "0710h", "0000h", "PUSH R1", "Помещение C001h в стек."),
    ("0008h", "0120h", "0021h", "MOV 0021h, R2", "Загрузка первого логического операнда."),
    ("000Ah", "0320h", "0022h", "OR R2, 0022h", "Формирование 0FFFh."),
    ("000Ch", "0420h", "0023h", "NOR R2, 0023h", "Формирование F000h."),
    ("000Eh", "0830h", "0000h", "POP R3", "Извлечение C001h из стека."),
    ("0010h", "0230h", "0030h", "MOV R3, 0030h", "Запись результата в ОЗУ."),
    ("0012h", "0140h", "0025h", "MOV 0025h, R4", "Загрузка FFFFh для получения нуля."),
    ("0014h", "0440h", "0025h", "NOR R4, 0025h", "Получение 0000h и установка Z=1."),
    ("0016h", "0A00h", "001Ah", "JZ 001Ah", "Условный переход должен выполниться."),
    ("0018h", "0150h", "0024h", "MOV 0024h, R5", "Команда должна быть пропущена."),
    ("001Ah", "0900h", "001Eh", "JMP 001Eh", "Безусловный переход к чтению результата DMA."),
    ("001Ch", "0160h", "0024h", "MOV 0024h, R6", "Команда должна быть пропущена."),
    ("001Eh", "0170h", "000Ah", "MOV 000Ah, R7", "Чтение первого слова, записанного КПДП."),
    ("0020h", "0000h", "0000h", "HLT", "Останов полной модели."),
]


RAM_ROWS = [
    ("000Ah", "0000h", "1111h", "Первое слово, переданное КПДП."),
    ("000Bh", "0000h", "2222h", "Второе слово, переданное КПДП."),
    ("000Ch", "0000h", "3333h", "Третье слово, переданное КПДП."),
    ("0020h", "8001h", "8001h", "Исходное отрицательное число для SRA."),
    ("0021h", "00F0h", "00F0h", "Первый логический операнд."),
    ("0022h", "0F0Fh", "0F0Fh", "Второй операнд для OR."),
    ("0023h", "0000h", "0000h", "Операнд для NOR."),
    ("0024h", "1234h", "1234h", "Контрольное значение пропущенных команд."),
    ("0025h", "FFFFh", "FFFFh", "Операнд для получения нулевого результата."),
    ("0030h", "0000h", "C001h", "Результат, записанный командой MOV R3, 0030h."),
]


SCHEME_ROWS = [
    ("А1", "A1_general_structure.drawio", "Общая структурная схема микро-ЭВМ."),
    ("А2", "A2_control_unit.drawio", "Устройство управления, дешифратор и фазовый автомат."),
    ("А3", "A3_datapath_alu_stack.drawio", "РОН, АЛУ, регистр флагов, стек и тракт записи результата."),
    ("А4", "A4_memory_dma_arbiter.drawio", "Память, кэш, КПДП, арбитр и предсказатель переходов A4."),
]


def hex_to_bin_word(value: str) -> str:
    raw = value.removesuffix("h")
    bits = f"{int(raw, 16):016b}"
    return " ".join(bits[i : i + 4] for i in range(0, 16, 4)) + "b"


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def style_cell(cell, font_size: int = 10, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), "80")
        node.set(qn("w:type"), "dxa")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(font_size)
            run.bold = bold


def set_table_font(table, font_size: int = 10) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            style_cell(cell, font_size=font_size, bold=(i == 0))


def set_cell_width(cell, width_cm: float) -> None:
    width = int(Cm(width_cm).twips)
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width_cm in zip(row.cells, widths_cm):
            set_cell_width(cell, width_cm)


def set_run_font(run, size: int = 14, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def normalize_document_fonts(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            if run.text:
                set_run_font(run, 14)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, 14, bold=True)
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 14)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 14)
    return paragraph


def add_listing_pair(doc: Document, addr: str, word0: str, word1: str, mnemonic: str) -> None:
    head = doc.add_paragraph()
    head.paragraph_format.left_indent = Cm(0.75)
    head.paragraph_format.space_before = Pt(3)
    head.paragraph_format.space_after = Pt(0)
    run = head.add_run(f"{addr}  {mnemonic}")
    set_run_font(run, 11, bold=True)

    detail = doc.add_paragraph()
    detail.paragraph_format.left_indent = Cm(1.25)
    detail.paragraph_format.space_after = Pt(2)
    text = (
        f"Word0: {word0} = {hex_to_bin_word(word0)}; "
        f"Word1: {word1} = {hex_to_bin_word(word1)}"
    )
    run = detail.add_run(text)
    set_run_font(run, 10)


def add_compact_line(doc: Document, title: str, detail: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{title} - ")
    set_run_font(run, 11, bold=True)
    run = paragraph.add_run(detail)
    set_run_font(run, 11)


def clean_cell_text(text: str) -> str:
    return " ".join(part.strip() for part in text.replace("\n", " ").split() if part.strip())


def insert_paragraph_before_table(table, text: str, size: int = 10, bold_prefix: str | None = None) -> None:
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.space_after = Pt(2)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, size, bold=True)
        run = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(run, size)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size)


def convert_tables_to_text(doc: Document) -> None:
    for table in list(doc.tables):
        rows = []
        for row in table.rows:
            values = [clean_cell_text(cell.text) for cell in row.cells]
            while values and values[-1] == "":
                values.pop()
            if values:
                rows.append(values)
        if not rows:
            table._tbl.getparent().remove(table._tbl)
            continue

        headers = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else []
        for row_index, values in enumerate(data_rows, start=1):
            parts = []
            for idx, value in enumerate(values):
                if not value:
                    continue
                header = headers[idx] if idx < len(headers) and headers[idx] else f"Поле {idx + 1}"
                if header == value:
                    parts.append(value)
                else:
                    parts.append(f"{header}: {value}")
            if not parts:
                continue
            prefix = f"{row_index}. "
            insert_paragraph_before_table(table, prefix + "; ".join(parts), size=10, bold_prefix=prefix)
        table._tbl.getparent().remove(table._tbl)


def add_program_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ А. Листинг тестовой программы", 1)
    add_body(
        doc,
        "В приложении приведён листинг программы в символьном, шестнадцатеричном и "
        "двоичном виде. Команда занимает два 16-разрядных слова: в первом слове "
        "размещаются код операции и номер регистра, во втором - адрес операнда "
        "или адрес перехода.",
    )
    for addr, w0, w1, mnemonic, _purpose in PROGRAM_ROWS:
        add_listing_pair(doc, addr, w0, w1, mnemonic)


def add_dump_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ Б. Контрольные дампы памяти", 1)
    add_body(
        doc,
        "Дампы используются для проверки функционального моделирования полной "
        "системы. ПЗУ должно содержать программу из приложения А. Для ОЗУ "
        "контролируются ячейки, которые читаются процессором, изменяются КПДП "
        "или записываются командой MOV R3, 0030h.",
    )
    for row_data in RAM_ROWS:
        addr, before, after, note = row_data
        add_compact_line(doc, addr, f"до запуска {before}; после HLT {after}; {note}")
    add_body(
        doc,
        "Для кэш-памяти в отчёт добавляется скриншот после моделирования: должны "
        "быть видны признаки hit/miss, действительные строки для адресов рабочих "
        "операндов и отсутствие dirty-бита, так как принята политика сквозной "
        "записи. Полные MIF-файлы подготовлены в каталоге memory проекта.",
    )


def add_schemes_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ В. Состав графической части", 1)
    add_body(
        doc,
        "Графическая часть проекта содержит четыре редактируемых листа формата "
        "drawio. Перед печатью их необходимо экспортировать в PDF или PNG, "
        "добавить рамку и основной штамп, если это требуется методическими "
        "указаниями кафедры.",
    )
    for row_data in SCHEME_ROWS:
        sheet, filename, note = row_data
        add_compact_line(doc, f"{sheet}: {filename}", note)
    add_body(
        doc,
        "В раздел функционального моделирования дополнительно вставляются реальные "
        "временные диаграммы АЛУ, стека, КПДП с арбитром и полной модели "
        "tb_system_variant4, снятые после запуска проекта в Quartus/ModelSim.",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(INPUT)

    # Keep page geometry close to the source report and leave appendices portrait.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT

    replacements = {
        "Разработка проекта велась с помощью программы Altera Quartus 9.1, которая позволяет разрабатывать схемы различной сложности, проверять их работоспособность путем подачи различных входных значений, а так же считывания выходных, а так же следить за тем, что происходит с разработанной схемой во время исполнения той или иной команды посредством временных диаграмм.": (
            "Проектные материалы подготовлены для реализации в САПР Altera Quartus версий 7.11-9.0, что соответствует требованиям задания. "
            "Функциональная проверка выполняется по VHDL-моделям и testbench; после переноса проекта в Quartus/ModelSim по тем же сигналам снимаются итоговые временные диаграммы и дампы памяти."
        ),
        "Принятый листинг удобно использовать и как программу, и как дамп ПЗУ. При необходимости двоичное представление слов получается прямым переводом соответствующих hex-кодов в 16-разрядную форму; например, слово 0110h соответствует коду 0000 0001 0001 0000b. Для инициализации ОЗУ данных используются значения, приведённые в таблице 15. Важно не путать одинаковые числовые адреса в таблицах 14 и 15: в первом случае они относятся к пространству команд, во втором - к пространству данных, что соответствует Гарвардской архитектуре.": (
            "Принятый листинг удобно использовать и как программу, и как дамп ПЗУ. Двоичное представление каждого 16-разрядного слова вынесено в приложение А, чтобы требование задания о символьном и бинарном листинге было выполнено явно. Для инициализации ОЗУ данных используются значения, приведённые в таблице 15. Важно не путать одинаковые числовые адреса в таблицах 14 и 15: в первом случае они относятся к пространству команд, во втором - к пространству данных, что соответствует Гарвардской архитектуре."
        ),
        "Таким образом, третий раздел задаёт полный набор исходных данных для моделирования и заранее фиксирует контрольные результаты, по которым удобно проверять работоспособность проекта. При фактическом оформлении курсовой работы в этот текст останется добавить только скриншоты временных диаграмм и, при необходимости, вынести полный листинг программы и расширенные дампы памяти в приложения. Содержательная часть проверки - что именно нужно наблюдать на диаграммах и какие значения считать правильными - уже определена.": (
            "Таким образом, третий раздел задаёт полный набор исходных данных для моделирования и заранее фиксирует контрольные результаты, по которым удобно проверять работоспособность проекта. В окончательный вариант записки необходимо добавить реальные скриншоты временных диаграмм и дампов памяти после запуска проекта в Quartus/ModelSim. Полный символьный и двоичный листинг программы, контрольные дампы и состав графической части приведены в приложениях."
        ),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])

    # Remove an orphaned figure from a previous draft in section 1.
    for idx, paragraph in enumerate(list(doc.paragraphs)):
        if paragraph.text.strip() == "Рис. 2.1.2 – условно-графическое обозначение блока памяти RAM":
            for victim in reversed(doc.paragraphs[max(0, idx - 2) : idx + 3]):
                delete_paragraph(victim)
            break

    # Replace the outdated 8-bit ALU fragment with variant-4 content.
    heading = None
    start_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Арифметико-логическое устройство":
            heading = paragraph
            start_idx = i
            break
    if heading is None:
        raise RuntimeError("ALU heading not found")

    end_idx = start_idx
    for i in range(start_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith("Блок регистров общего назначения"):
            end_idx = i
            break
    for victim in reversed(doc.paragraphs[start_idx + 1 : end_idx]):
        delete_paragraph(victim)

    set_paragraph_text(heading, "2.3 Арифметико-логическое устройство и блок РОН")
    heading.style = "Heading 2"
    p = heading
    inserted = [
        "Вычислительная часть микро-ЭВМ включает 16-разрядное арифметико-логическое устройство, регистровый файл 12 x 16 и регистр флагов FR. Вариант 4 требует реализации операций INCS, OR, NOR и SRA; операции ADDC, AND, NOT, ROL и CMP к данному варианту не относятся.",
        "На вход A АЛУ поступает содержимое выбранного регистра Rd. На вход B подаётся либо второй операнд из памяти/регистра, либо расширенное до 16 бит значение флага S для операции INCS. Код ALU_OP выбирает один из режимов: OR, NOR, SRA, INCS или PASS для служебной передачи данных.",
        "После выполнения операции результат Y[15:0] возвращается в регистровый файл, а признаки Z, S, C и O записываются в регистр флагов. Для SRA старший бит сохраняется как знаковый, младший вытесненный бит записывается во флаг C. Для INCS к операнду прибавляется текущее значение флага S, сохранённое до начала операции.",
        "Основные сигналы АЛУ: A[15:0], B[15:0], ALU_OP[2:0], Y[15:0], Z, S, C, O. Основные сигналы блока РОН: RF_RA1, RF_RA2, RF_WA, RF_WE, RF_DIN[15:0], RF_DOUT1[15:0], RF_DOUT2[15:0].",
    ]
    for text in inserted:
        p = insert_paragraph_after(p, text, style="Normal")
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_run_font(run, 14)

    add_program_appendix(doc)
    add_dump_appendix(doc)
    add_schemes_appendix(doc)
    convert_tables_to_text(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
